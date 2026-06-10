# -*- coding: utf-8 -*-
"""F13: 最终验证 PLUS 修订版 v3 docx 的完整性。"""
from docx import Document
import os

V3 = r'E:\大学\万物春\erci\郑春铃+横州市土地利用变化与生态安全评价(1)_PLUS修订版_v3.docx'
ORIG = r'E:\大学\万物春\erci\郑春铃+横州市土地利用变化与生态安全评价(1).docx.bak_before_plus_revision'

print('==== 最终修订版 v3 统计 ====')
print(f'文件: {V3}')
print(f'大小: {os.path.getsize(V3):,} bytes')

doc = Document(V3)
paras = doc.paragraphs
print(f'段落数: {len(paras)}  (原 470, 新增至少 19)')
print(f'表格数: {len(doc.tables)}  (原 16, +表14)')

# 图片关系
img_count = 0
for rel in doc.part.rels.values():
    if 'image' in rel.target_ref:
        img_count += 1
print(f'图片关系数: {img_count}  (原 19, +图1\u2032/图3\u2032/图12\u2032/图16 = +4)')

all_text = '\n'.join(p.text for p in paras)

print('\n==== 13 项细节修复验证 ====')

checks = [
    ('F1+F3 PLUS 三段', '6 个驱动因子', '应展示真实因子数'),
    ('F1+F3 政策约束', '未额外引入耕地保护红线', '应说明未引入政策'),
    ('F1+F3 训练参数', 'n_estimators=30', '应给出 RF 参数'),
    ('F2 未利用地训练失败', '未利用地因 2010—2015 年扩展样本不足', '应披露训练失败'),
    ('F2 CARS 简化披露', '本研究的 CARS 简化实现', '应披露 CARS 简化'),
    ('F4 展望-PLUS局限', 'PLUS 模型情景模拟方面存在两点局限', '应在展望中说明 PLUS 局限'),
    ('F4 展望-baseline', '与传统 CA-Markov、FLUS 等模型进行 baseline 对比', '应在展望中提 baseline'),
    ('F4 展望-外部校验', '待权威机构发布 2025 年实测', '应在展望中提外部校验'),
    ('F5 图1\u2032 修订版', '图1\u2032 横州市研究技术路线图', '应有修订版技术路线图标题'),
    ('F6 图3\u2032 修订版', '图3\u2032 横州市 2010—2025', '应有修订版图3标题'),
    ('F7 图12\u2032 修订版', '图12\u2032 横州市生态安全指数', '应有修订版图12标题'),
    ('F8 图16 驱动因子', '图16 横州市 PLUS 模型驱动因子', '应有图16标题'),
    ('F9 表14 景观格局', '表14 横州市 PLUS 模型主要地类景观格局', '应有表14标题'),
    ('F10 表内 2025（模拟）', '2025（模拟）', '表格内 2025 应标注'),
    ('F11 中文摘要 PLUS 全称', 'PLUS（Patch-generating Land Use Simulation）', '中文摘要应全称'),
    ('F11 英文摘要 PLUS 全称', 'PLUS (Patch-generating Land Use Simulation)', '英文摘要应全称'),
    ('F12 修正 [10] 文献', '李德仁, 张良培, 夏桂松', '应替换为真实文献'),
    ('F12 删除虚假"刘桂霞 2025"', '刘桂霞', '应删除（仅检查不应在文献列表）'),
]

passed, failed = 0, 0
for label, kw, expected in checks:
    n = all_text.count(kw)
    if 'F12 删除' in label:
        # 仅检查不应出现在文献列表中
        ok = n == 0
    else:
        ok = n > 0
    flag = 'PASS' if ok else 'FAIL'
    if ok:
        passed += 1
    else:
        failed += 1
    print(f'  [{flag:4}] {label}: "{kw}" → 命中 {n} 处')

print(f'\n通过: {passed}/{len(checks)}, 失败: {failed}/{len(checks)}')

# 关键旧表述彻底清除验证
print('\n==== 关键旧虚假/错误表述清除 ====')
bad_checks = [
    '基于2010、2015、2020和2025年四期土地利用遥感解译数据，采用土地利用动态度法',
    '来源于中国科学院资源环境科学与数据中心。数据经过人工目视解译',
    '2025年的约50%，城镇化进程持续推进',
    '刘桂霞. 遥感与GIS融合技术',  # 虚假文献
    '总体精度（Overall Accuracy）达到87.3%',  # 占位精度
    'Kappa系数为0.82',
    '采用Markov链模型基于2015—2020年实测的土地利用转移概率矩阵，外推至2025年得到各地类的数量需求。考虑到2020年至研究终点2025年间的政策约束',  # 政策约束占位
]
all_clear = True
for bad in bad_checks:
    n = all_text.count(bad)
    flag = 'OK' if n == 0 else 'FAIL'
    if n != 0:
        all_clear = False
    print(f'  [{flag:4}] "{bad[:50]}..." → {n} 处')
print(f'\n旧表述清除: {"全部清除" if all_clear else "仍有残留!"}')

# 新增小节结构检查
print('\n==== PLUS 节结构 ====')
plus_start = -1
for i, p in enumerate(paras):
    if '（三）PLUS模型情景模拟' in p.text:
        plus_start = i
        break
if plus_start > 0:
    print(f'PLUS 节起点: P{plus_start}')
    for i in range(plus_start, min(plus_start + 15, len(paras))):
        style = paras[i].style.name if paras[i].style else ''
        text = paras[i].text[:60]
        print(f'  P{i}|{style:14} {text}')

# 第六章展望段
print('\n==== 第六章展望段 ====')
for i, p in enumerate(paras):
    if '本研究存在以下不足' in p.text or '未来研究可从以下方面深化' in p.text:
        style = p.style.name if p.style else ''
        print(f'P{i}|{style}: {p.text[:200]}')

print('\n==== 完成 ====')
