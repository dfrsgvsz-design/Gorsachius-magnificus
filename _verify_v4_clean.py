# -*- coding: utf-8 -*-
"""验证 v4 全文 OMML 公式分布，确认无残留污染"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.oxml.ns import qn
import zipfile

V4 = r'E:\大学\万物春\erci\郑春铃+横州市土地利用变化与生态安全评价(1)_PLUS修订版_v4.docx'

doc = Document(V4)

print('==== v4 全文 OMML 公式分布 ====')
omml_paras = []
for i, p in enumerate(doc.paragraphs):
    n = len(p._p.findall('.//' + qn('m:oMath'))) + len(p._p.findall('.//' + qn('m:oMathPara')))
    if n > 0:
        omml_paras.append((i, n, p.text[:80]))

print(f'共 {len(omml_paras)} 个段落含公式（应仅在公式正文/变量说明段）')
for i, n, t in omml_paras:
    print(f'  P{i}: {n} 公式 | {t}')

# 检查 PLUS 节区域（P193-P204）应该 0 公式
print('\n==== PLUS 节（P193-P204）应全 0 公式 ====')
plus_clean = True
for i in range(193, 205):
    n = len(doc.paragraphs[i]._p.findall('.//' + qn('m:oMath'))) + len(doc.paragraphs[i]._p.findall('.//' + qn('m:oMathPara')))
    if n > 0:
        plus_clean = False
        print(f'  [FAIL] P{i}: {n} 公式残留')
if plus_clean:
    print('  [PASS] PLUS 节全部干净')

# 检查 XML 中是否有任何 LaTeX 源码字符串
print('\n==== 全文 XML 字符串检查 ====')
with zipfile.ZipFile(V4) as z:
    with z.open('word/document.xml') as f:
        xml = f.read().decode('utf-8')

bad = ['x_{ij}', '\\prime', '\\frac', '\\min', '\\max']
for s in bad:
    n = xml.count(s)
    flag = 'OK' if n == 0 else 'FAIL'
    print(f'  [{flag}] "{s}": {n} 处')

# 显示 P193 附近 XML 净化后的结构
print('\n==== P193 净化后 XML 片段 ====')
idx = xml.find('（三）PLUS模型情景模拟')
if idx > 0:
    snippet = xml[max(0, idx-200):idx+400]
    print(snippet)

# 文件大小、表格、图片
import os
print(f'\n==== v4 总览 ====')
print(f'  文件大小: {os.path.getsize(V4):,} bytes')
print(f'  段落数: {len(doc.paragraphs)}')
print(f'  表格数: {len(doc.tables)}')
img_count = sum(1 for r in doc.part.rels.values() if 'image' in r.target_ref)
print(f'  图片数: {img_count}')
