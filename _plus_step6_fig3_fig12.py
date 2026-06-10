# -*- coding: utf-8 -*-
"""Step 6: 重画图3（四期土地利用现状）+ 图12（ESI 折线图），含 2025 PLUS 标注。
然后插入到 docx 中相应位置后面。
"""
import os
import json
import numpy as np
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from matplotlib import rcParams
from matplotlib.colors import ListedColormap
from matplotlib.patches import Patch
import rasterio

rcParams['font.sans-serif'] = ['Microsoft YaHei', 'SimHei', 'DejaVu Sans']
rcParams['axes.unicode_minus'] = False
rcParams['font.size'] = 10

WORK = r'E:\大学\万物春\erci\_plus_workspace'

# 7 类土地利用配色（基于 IGBP 风格 + 中国典型用色）
CLASS_COLORS = {
    1: '#FFE5B4',  # 耕地 - 浅米色
    2: '#228B22',  # 林地 - 深绿
    3: '#90EE90',  # 灌木 - 浅绿
    4: '#FAFAD2',  # 草地 - 浅黄
    5: '#4682B4',  # 水域 - 蓝
    7: '#A0A0A0',  # 未利用地 - 灰
    8: '#DC143C',  # 建设用地 - 红
}
CLASS_NAMES = {1: '耕地', 2: '林地', 3: '灌木', 4: '草地',
               5: '水域', 7: '未利用地', 8: '建设用地'}
CLASSES = [1, 2, 3, 4, 5, 7, 8]


# ============================================================
# F6: 图3 四期土地利用现状图
# ============================================================
def make_fig3():
    fig, axes = plt.subplots(2, 2, figsize=(14, 12), dpi=200)
    axes = axes.flatten()

    years = [2010, 2015, 2020, 2025]
    titles = ['(a) 2010 年', '(b) 2015 年', '(c) 2020 年', '(d) 2025 年（PLUS 模拟情景）']

    # 构造颜色映射
    cmap_list = [CLASS_COLORS[c] for c in CLASSES]
    cmap = ListedColormap(cmap_list)

    for ax, year, title in zip(axes, years, titles):
        path = os.path.join(WORK, f'lu_{year}.tif')
        with rasterio.open(path) as src:
            arr = src.read(1)
            nd = src.nodata

        # 下采样
        ds = 4
        arr_ds = arr[::ds, ::ds]
        # 映射 gridcode 到 0-6 索引
        mapped = np.full(arr_ds.shape, np.nan, dtype=float)
        for i, c in enumerate(CLASSES):
            mapped[arr_ds == c] = i
        mapped[arr_ds == nd] = np.nan

        im = ax.imshow(mapped, cmap=cmap, vmin=-0.5, vmax=len(CLASSES) - 0.5,
                       interpolation='nearest', aspect='equal')
        ax.set_title(title, fontsize=12, fontweight='bold')
        ax.axis('off')

        # 在 2025 子图右上角加 PLUS 标注框
        if year == 2025:
            ax.text(0.97, 0.97, 'PLUS\n模拟', transform=ax.transAxes,
                    ha='right', va='top', fontsize=10, fontweight='bold',
                    color='#D97706',
                    bbox=dict(boxstyle='round,pad=0.4', facecolor='#FEF3C7',
                              edgecolor='#D97706', linewidth=1.5))

    # 图例
    handles = [Patch(facecolor=CLASS_COLORS[c], edgecolor='black',
                     label=CLASS_NAMES[c]) for c in CLASSES]
    fig.legend(handles=handles, loc='lower center', ncol=7, fontsize=11,
               frameon=True, bbox_to_anchor=(0.5, -0.005))

    plt.tight_layout()
    plt.subplots_adjust(bottom=0.08)
    out_path = os.path.join(WORK, 'fig3_landuse_4periods_revised.png')
    plt.savefig(out_path, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close()
    print(f'  [OK] {out_path}')
    return out_path


# ============================================================
# F7: 图12 ESI 折线图
# ============================================================
def make_fig12():
    years = [2010, 2015, 2020, 2025]
    esi = [0.718, 0.460, 0.390, 0.283]
    levels = ['Ⅱ 较安全', 'Ⅲ 预警', 'Ⅳ 中度不安全', 'Ⅳ 中度不安全']

    fig, ax = plt.subplots(figsize=(10, 6), dpi=200)

    # 实测期（2010-2020）+ 模拟期（2020-2025）分段画
    ax.plot(years[:3], esi[:3], 'o-', color='#1E40AF', linewidth=2.5,
            markersize=12, markerfacecolor='#3B82F6', markeredgecolor='#1E40AF',
            markeredgewidth=2, label='实测情景 (2010—2020)', zorder=3)
    # 2020->2025 模拟段（虚线）
    ax.plot(years[2:], esi[2:], '--', color='#F59E0B', linewidth=2.5, zorder=2)
    # 2025 点：空心标记
    ax.plot(2025, 0.283, 'o', markersize=14, markerfacecolor='white',
            markeredgecolor='#D97706', markeredgewidth=3,
            label='2025 PLUS 模拟情景', zorder=4)

    # 数值标注
    for x, y, lvl in zip(years, esi, levels):
        ax.annotate(f'ESI={y:.3f}\n({lvl})',
                    xy=(x, y), xytext=(0, 18), textcoords='offset points',
                    ha='center', fontsize=10,
                    bbox=dict(boxstyle='round,pad=0.3',
                              facecolor='white' if x != 2025 else '#FEF3C7',
                              edgecolor='#D97706' if x == 2025 else '#94A3B8',
                              linewidth=1))

    # 等级分界线（横向）
    level_thresholds = {0.8: 'Ⅰ 安全', 0.6: 'Ⅱ 较安全', 0.4: 'Ⅲ 预警', 0.2: 'Ⅳ 中度不安全'}
    for thr, name in level_thresholds.items():
        ax.axhline(y=thr, color='gray', linestyle=':', linewidth=0.8, alpha=0.5)
        ax.text(2009.5, thr + 0.01, name, fontsize=8, color='gray', alpha=0.7)

    # 模拟期阴影
    ax.axvspan(2020, 2025, alpha=0.08, color='#F59E0B', label='PLUS 模拟期')

    ax.set_xlabel('年份', fontsize=12, fontweight='bold')
    ax.set_ylabel('生态安全指数 ESI', fontsize=12, fontweight='bold')
    ax.set_title('图12 横州市生态安全指数演变趋势（含 PLUS 模拟期）',
                 fontsize=13, fontweight='bold')
    ax.set_xticks(years)
    ax.set_xlim(2008, 2027)
    ax.set_ylim(0.15, 0.85)
    ax.grid(True, alpha=0.3)
    ax.legend(loc='upper right', fontsize=10, framealpha=0.95)

    plt.tight_layout()
    out_path = os.path.join(WORK, 'fig12_esi_revised.png')
    plt.savefig(out_path, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close()
    print(f'  [OK] {out_path}')
    return out_path


# ============================================================
# 插入到 docx
# ============================================================
def insert_to_docx():
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from copy import deepcopy

    DOC_PATH = r'E:\大学\万物春\erci\郑春铃+横州市土地利用变化与生态安全评价(1)_PLUS修订版_v2.docx'
    DST = DOC_PATH.replace('_v2.docx', '_v3.docx')

    doc = Document(DOC_PATH)
    paras = doc.paragraphs
    print(f'\n读取 v2: {DOC_PATH}')
    print(f'  段落数: {len(paras)}, 表格数: {len(doc.tables)}')

    def insert_after(template_para, text=''):
        new_p = deepcopy(template_para._p)
        for t in new_p.findall('.//' + qn('w:t')):
            t.text = ''
        # ★ 关键修复：清空所有 OMML 公式元素，避免污染新段 ★
        for omath in list(new_p.findall('.//' + qn('m:oMath'))):
            parent = omath.getparent()
            if parent is not None:
                parent.remove(omath)
        for omp in list(new_p.findall('.//' + qn('m:oMathPara'))):
            parent = omp.getparent()
            if parent is not None:
                parent.remove(omp)
        runs_xml = new_p.findall(qn('w:r'))
        if runs_xml:
            for r in runs_xml[1:]:
                new_p.remove(r)
            first_t = new_p.find('.//' + qn('w:t'))
            if first_t is None:
                from docx.oxml import OxmlElement
                r0 = new_p.find(qn('w:r'))
                t_new = OxmlElement('w:t')
                t_new.set(qn('xml:space'), 'preserve')
                t_new.text = text
                r0.append(t_new)
            else:
                first_t.text = text
        template_para._p.addnext(new_p)
        from docx.text.paragraph import Paragraph
        return Paragraph(new_p, template_para._parent)

    # 图3 标题段：'图3 横州市2010、2015、2020、2025年四期土地利用现状图（2025年为PLUS模拟情景）'
    fig3_idx = -1
    for i, p in enumerate(paras):
        if '图3' in p.text and '四期土地利用现状图' in p.text:
            fig3_idx = i
            break
    if fig3_idx >= 0:
        anchor = paras[fig3_idx]
        print(f'  [F6] 图3 锚点: P{fig3_idx}')
        img_para = insert_after(anchor, '')
        run = img_para.add_run()
        run.add_picture(os.path.join(WORK, 'fig3_landuse_4periods_revised.png'),
                        width=Inches(6.5))
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = insert_after(img_para,
            '图3\u2032 横州市 2010—2025 年四期土地利用现状图（修订版，2025 年子图明确标注为 PLUS 模拟情景）')
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        print('  [OK] 图3\u2032 已插入')

    # 图12 标题段
    paras = doc.paragraphs
    fig12_idx = -1
    for i, p in enumerate(paras):
        if '图12' in p.text and '生态安全指数' in p.text:
            fig12_idx = i
            break
    if fig12_idx >= 0:
        anchor = paras[fig12_idx]
        print(f'  [F7] 图12 锚点: P{fig12_idx}')
        img_para = insert_after(anchor, '')
        run = img_para.add_run()
        run.add_picture(os.path.join(WORK, 'fig12_esi_revised.png'),
                        width=Inches(6.0))
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = insert_after(img_para,
            '图12\u2032 横州市生态安全指数演变趋势（修订版，2010—2020 年为实测段，2020—2025 年为 PLUS 模拟段，2025 年标记为空心圆）')
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        print('  [OK] 图12\u2032 已插入')

    doc.save(DST)
    print(f'\n保存至: {DST}')
    print(f'最终: 段落 {len(doc.paragraphs)} | 表格 {len(doc.tables)}')
    return DST


if __name__ == '__main__':
    print('==== F6: 重画图3 ====')
    make_fig3()
    print('\n==== F7: 重画图12 ====')
    make_fig12()
    print('\n==== 插入到 docx ====')
    insert_to_docx()
    print('\nDone.')
