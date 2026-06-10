# -*- coding: utf-8 -*-
"""读取论文撰写规范 docx 全文内容"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document

SPEC = r'E:\大学\万物春\erci\附件三：论文撰写规范 (1).docx'
doc = Document(SPEC)

print(f'==== 总览 ====')
print(f'段落数: {len(doc.paragraphs)}')
print(f'表格数: {len(doc.tables)}')
print()
print('==== 全文段落 ====')
for i, p in enumerate(doc.paragraphs):
    style = p.style.name if p.style else 'None'
    text = p.text
    if not text.strip() and 'Heading' not in style:
        continue
    print(f'[P{i:3d}|{style:15s}] {text}')

print()
print('==== 表格 ====')
for ti, tbl in enumerate(doc.tables):
    print(f'\n--- Table {ti} ({len(tbl.rows)} rows × {len(tbl.columns)} cols) ---')
    for ri, row in enumerate(tbl.rows):
        for ci, cell in enumerate(row.cells):
            txt = cell.text.replace('\n', ' | ')
            print(f'  R{ri}C{ci}: {txt[:120]}')
