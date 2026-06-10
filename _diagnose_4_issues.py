# -*- coding: utf-8 -*-
"""诊断 4 个新发现问题"""
import sys, io, re
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.oxml.ns import qn

DOC = r'E:\大学\万物春\erci\郑春铃+横州市土地利用变化与生态安全评价(1)_PLUS修订版_最终_格式合规.docx'
doc = Document(DOC)
paras = doc.paragraphs
all_text = '\n'.join(p.text for p in paras)

print('=' * 70)
print('问题 1：参考文献引用顺序')
print('=' * 70)
# 找出所有 [N] 在正文中的出现位置（按段落序）
citation_order = []  # [(段号, 编号)]
ref_start = -1
for i, p in enumerate(paras):
    if p.text.strip() == '参考文献':
        ref_start = i
        break

# 在正文范围（参考文献节之前）扫描 [N] 引用
for i in range(ref_start):
    t = paras[i].text
    for m in re.finditer(r'\[(\d+)\]', t):
        n = int(m.group(1))
        citation_order.append((i, n))

# 提取首次引用顺序
seen = set()
first_cite_order = []
for pi, n in citation_order:
    if n not in seen:
        first_cite_order.append((pi, n))
        seen.add(n)

print(f'正文中引用的不同编号数: {len(first_cite_order)}')
print(f'正文中首次引用顺序: {[n for _, n in first_cite_order[:15]]}...')
print(f'当前参考文献列表是 1—{len(first_cite_order)}')

# 检查是否按引用序排
ideal_order = list(range(1, len(first_cite_order) + 1))
actual_order = [n for _, n in first_cite_order]
if actual_order == ideal_order:
    print('✅ 引用顺序与列表顺序一致')
else:
    print('❌ 引用顺序与列表不一致')
    print(f'   理想: {ideal_order[:15]}')
    print(f'   实际: {actual_order[:15]}')

print()
print('=' * 70)
print('问题 2：图表编号乱序')
print('=' * 70)
print('--- 所有图标题 ---')
fig_caps = []
for i, p in enumerate(paras):
    t = p.text.strip()
    m = re.match(r'^图\s*(\d+)([\u2032\u2019\u2018\']?)\s*(.*)', t)
    if m:
        fig_caps.append((i, m.group(1), m.group(2), m.group(3)[:50]))
        print(f'  P{i}: 图{m.group(1)}{m.group(2)} {m.group(3)[:60]}')

print('\n--- 所有表标题 ---')
tbl_caps = []
for i, p in enumerate(paras):
    t = p.text.strip()
    m = re.match(r'^表\s*(\d+)([\u2032\u2019\u2018\']?)\s*(.*)', t)
    if m:
        tbl_caps.append((i, m.group(1), m.group(2), m.group(3)[:50]))
        print(f'  P{i}: 表{m.group(1)}{m.group(2)} {m.group(3)[:60]}')

print()
print('=' * 70)
print('问题 3：表格内首行缩进 2 字符')
print('=' * 70)
indent_count = 0
for ti, tbl in enumerate(doc.tables):
    for ri, row in enumerate(tbl.rows):
        for ci, cell in enumerate(row.cells):
            for p in cell.paragraphs:
                pf = p.paragraph_format
                if pf.first_line_indent and pf.first_line_indent.cm > 0.3:
                    indent_count += 1
                    if indent_count <= 10:
                        print(f'  Table{ti} R{ri}C{ci}: first_line_indent={pf.first_line_indent.cm:.2f} cm | text="{p.text[:40]}"')
print(f'\n表格内含首行缩进 >0.3cm 的段落数: {indent_count}')

print()
print('=' * 70)
print('问题 4：字体乱入位置（统计各 eastAsia 字体分布）')
print('=' * 70)
font_dist = {}
fontsize_dist = {}
for i, p in enumerate(paras):
    if not p.runs:
        continue
    for run in p.runs:
        if not run.text.strip():
            continue
        rPr = run._element.find(qn('w:rPr'))
        ea = None
        if rPr is not None:
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is not None:
                ea = rFonts.get(qn('w:eastAsia'))
        ea = ea or '(继承)'
        font_dist[ea] = font_dist.get(ea, 0) + 1
        sz = run.font.size
        if sz:
            sz_pt = sz.pt
            fontsize_dist[sz_pt] = fontsize_dist.get(sz_pt, 0) + 1
print('eastAsia 字体分布:')
for fn, cnt in sorted(font_dist.items(), key=lambda x: -x[1]):
    print(f'  {fn:20s}: {cnt} 个 run')
print('\n字号分布 (pt):')
for sz, cnt in sorted(fontsize_dist.items()):
    print(f'  {sz:>6.1f} pt: {cnt} 个 run')

# 找出特殊段：标题/图表标题被改为宋体的（应该保留原样式）
print('\n--- 图/表标题段的当前字体（应小四号宋体） ---')
for i, p in enumerate(paras):
    t = p.text.strip()
    if re.match(r'^(图|表)\s*\d+', t):
        if p.runs:
            run = p.runs[0]
            sz = run.font.size
            rPr = run._element.find(qn('w:rPr'))
            ea = None
            if rPr is not None:
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is not None:
                    ea = rFonts.get(qn('w:eastAsia'))
            sz_str = f'{sz.pt}pt' if sz else '继承'
            print(f'  P{i}: {t[:25]} → size={sz_str}, font={ea}')
