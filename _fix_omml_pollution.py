# -*- coding: utf-8 -*-
"""修复 PLUS 节段落中误嵌的 OMML 公式（从 P192 "式中：..."模板段继承而来）

污染范围：P193 (PLUS 节起点) 到 P204 (表14 标题) 共 12 个段落，每段含 7 个误嵌公式。
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.oxml.ns import qn
from copy import deepcopy
import shutil

V3 = r'E:\大学\万物春\erci\郑春铃+横州市土地利用变化与生态安全评价(1)_PLUS修订版_v3.docx'
V4 = V3.replace('_v3.docx', '_v4.docx')

# 备份
shutil.copy(V3, V4)
print(f'已复制 v3 -> v4: {V4}')

doc = Document(V4)
paras = doc.paragraphs
print(f'总段落: {len(paras)}')

# 识别 PLUS 节段落范围
plus_start, plus_end = -1, -1
for i, p in enumerate(paras):
    if '（三）PLUS模型情景模拟' in p.text:
        plus_start = i
    elif plus_start > 0 and '表14 横州市 PLUS 模型主要地类景观格局' in p.text:
        plus_end = i
        break

if plus_start < 0 or plus_end < 0:
    print(f'[ERROR] PLUS 节定位失败 start={plus_start} end={plus_end}')
    sys.exit(1)

print(f'PLUS 节范围: P{plus_start} 到 P{plus_end} (共 {plus_end - plus_start + 1} 段)')

# 清理目标段落中的 m:oMath / m:oMathPara 元素
M_OMATH = qn('m:oMath')
M_OMATH_PARA = qn('m:oMathPara')

total_removed = 0
for i in range(plus_start, plus_end + 1):
    p = paras[i]
    pElem = p._p
    text_snip = p.text[:50] if p.text else '(空)'

    # 找所有 OMML 元素（注意 .// 是递归）
    omath_list = pElem.findall('.//' + M_OMATH)
    omath_para_list = pElem.findall('.//' + M_OMATH_PARA)
    all_omml = omath_list + omath_para_list

    if not all_omml:
        print(f'  P{i}: 无公式 | text="{text_snip}"')
        continue

    # 删除所有公式元素
    for omml in all_omml:
        # 获取父节点并 remove
        parent = omml.getparent()
        if parent is not None:
            parent.remove(omml)
            total_removed += 1

    print(f'  P{i}: 删除 {len(all_omml)} 个公式 | text="{text_snip}"')

print(f'\n总删除: {total_removed} 个误嵌 OMML 公式')

doc.save(V4)
print(f'保存至: {V4}')

# 验证
print('\n==== 验证 v4 ====')
doc2 = Document(V4)
for i in range(plus_start, plus_end + 1):
    p = doc2.paragraphs[i]
    omml_count = len(p._p.findall('.//' + M_OMATH)) + len(p._p.findall('.//' + M_OMATH_PARA))
    if omml_count > 0:
        print(f'  [仍有] P{i}: {omml_count} 个公式')
    else:
        print(f'  [OK]   P{i}: 0 公式')

# 验证 P189-P192 真实公式未受影响
print('\n==== 验证 P189-P192 真实公式未受影响 ====')
for i in range(189, 193):
    if i < len(doc2.paragraphs):
        p = doc2.paragraphs[i]
        omml_count = len(p._p.findall('.//' + M_OMATH)) + len(p._p.findall('.//' + M_OMATH_PARA))
        print(f'  P{i}: {omml_count} 公式 | text="{p.text[:50]}"')
