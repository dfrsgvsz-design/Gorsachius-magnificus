# -*- coding: utf-8 -*-
"""验证 PLUS 修订版 docx 的完整性。"""
from docx import Document

SRC = r'E:\大学\万物春\erci\郑春铃+横州市土地利用变化与生态安全评价(1)_PLUS修订版.docx'
ORIG = r'E:\大学\万物春\erci\郑春铃+横州市土地利用变化与生态安全评价(1).docx.bak_before_plus_revision'

print('==== 修订版统计 ====')
doc = Document(SRC)
paras = doc.paragraphs
print(f'段落数: {len(paras)}  (原 470, 期望 480)')
print(f'表格数: {len(doc.tables)}  (原 16)')

# 计算图片数（关系中的图片）
from docx.oxml.ns import qn
img_count = 0
for rel in doc.part.rels.values():
    if 'image' in rel.target_ref:
        img_count += 1
print(f'图片关系数: {img_count}')

all_text = '\n'.join(p.text for p in paras)

print('\n==== 关键关键词命中数 ====')
checks = [
    ('PLUS', '出现次数'),
    ('Patch-generating Land Use Simulation', 'PLUS全称'),
    ('模拟情景', '"模拟情景"标签'),
    ('（三）PLUS模型情景模拟', '新增小节标题'),
    ('1. PLUS模型原理', '新增子节1'),
    ('2. 驱动因子选择', '新增子节2'),
    ('3. 土地利用需求预测', '新增子节3'),
    ('4. 模型精度验证', '新增子节4'),
    ('Markov', 'Markov链'),
    ('Kappa', 'Kappa精度'),
    ('Liang X', 'PLUS参考文献'),
    ('CNLUCC', 'CNLUCC数据'),
    ('CLCD', 'CLCD数据'),
    ('2025年（PLUS模拟情景）', '2025标注'),
]
for kw, label in checks:
    n = all_text.count(kw)
    flag = 'OK' if n > 0 else 'MISS'
    print(f'  [{flag:4}] {label} "{kw}": {n} 处')

print('\n==== 旧表述残留检查（应为 0） ====')
bad_checks = [
    '基于2010、2015、2020和2025年四期土地利用遥感解译数据，采用土地利用动态度法',
    '来源于中国科学院资源环境科学与数据中心。数据经过人工目视解译',
    '2025年的约50%，城镇化进程持续推进',
]
ok_count = 0
for bad in bad_checks:
    n = all_text.count(bad)
    flag = 'OK' if n == 0 else 'FAIL'
    print(f'  [{flag:4}] "{bad[:40]}...": {n} 处')
    if n == 0:
        ok_count += 1
print(f'\n旧表述清除: {ok_count}/{len(bad_checks)}')

print('\n==== 第二章末尾结构（验证新插入小节）====')
# 找 (三) PLUS 小节，输出前后段落
for i, p in enumerate(paras):
    if '（三）PLUS' in p.text:
        for j in range(max(0, i - 2), min(len(paras), i + 12)):
            style = paras[j].style.name if paras[j].style else ''
            text = paras[j].text[:80]
            print(f'  P{j}|{style:14} {text}')
        break

print('\n==== 表格内容快速检查（2025 数据应保留）====')
table_check = [
    (2, 0, '土地 | 类型'),   # 表2 表头
    (8, 22, '2025'),         # 表8 R22
    (15, 4, '2025'),         # 表13/15 R4
]
for ti, ri, expected in table_check:
    if ti < len(doc.tables) and ri < len(doc.tables[ti].rows):
        cells = [c.text.strip() for c in doc.tables[ti].rows[ri].cells]
        cell0 = cells[0]
        flag = 'OK' if expected in cell0 else 'MISS'
        print(f'  [{flag:4}] Table{ti} R{ri}[0]="{cell0[:40]}" (期望含"{expected}")')

print('\n==== 完成 ====')
