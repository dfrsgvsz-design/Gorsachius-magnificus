# -*- coding: utf-8 -*-
"""v2 最终验证：图表编号连续性 + 表格首行缩进 + 字体一致性"""
import sys, io, re
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.oxml.ns import qn

DST = r'E:\大学\万物春\erci\郑春铃+横州市土地利用变化与生态安全评价(1)_PLUS修订版_最终_格式合规_v2.docx'
doc = Document(DST)
paras = doc.paragraphs

print(f'==== v2 总览 ====')
print(f'段落数: {len(paras)} (原 489, 删 9, 应 480)')
print(f'表格数: {len(doc.tables)}')

# 验证图标题连续性
print('\n==== 图标题（按物理顺序）====')
fig_nums = []
for i, p in enumerate(paras):
    t = p.text.strip()
    m = re.match(r'^图\s*(\d+)([\u2032\u2019\']?)\s', t)
    if m:
        n = int(m.group(1))
        prime = m.group(2)
        fig_nums.append(n)
        flag = '✗ 撇号未清除' if prime else '✓'
        print(f'  P{i}: 图{n}{prime}  {t[:50]} {flag}')

if fig_nums == list(range(1, len(fig_nums) + 1)):
    print(f'\n✅ 图编号连续：1—{len(fig_nums)}')
else:
    print(f'\n❌ 图编号不连续: {fig_nums}')

print('\n==== 表标题（按物理顺序）====')
tbl_nums = []
for i, p in enumerate(paras):
    t = p.text.strip()
    m = re.match(r'^表\s*(\d+)([\u2032\u2019\']?)\s', t)
    if m:
        n = int(m.group(1))
        prime = m.group(2)
        tbl_nums.append(n)
        flag = '✗ 撇号未清除' if prime else '✓'
        print(f'  P{i}: 表{n}{prime}  {t[:50]} {flag}')

if tbl_nums == list(range(1, len(tbl_nums) + 1)):
    print(f'\n✅ 表编号连续：1—{len(tbl_nums)}')
else:
    print(f'\n❌ 表编号不连续: {tbl_nums}')

# 表格首行缩进
print('\n==== 表格内首行缩进检查 ====')
indent_count = 0
for ti, tbl in enumerate(doc.tables):
    for ri, row in enumerate(tbl.rows):
        for ci, cell in enumerate(row.cells):
            for p in cell.paragraphs:
                pf = p.paragraph_format
                if pf.first_line_indent and pf.first_line_indent.cm > 0.1:
                    indent_count += 1
print(f'表格内仍有首行缩进 >0.1cm 的段落: {indent_count} 个')
if indent_count == 0:
    print('✅ 表格内首行缩进已全部清除')

# 字号字体分布
print('\n==== 字体分布 ====')
font_dist = {}
size_dist = {}
for p in paras:
    for run in p.runs:
        if not run.text.strip():
            continue
        rPr = run._element.find(qn('w:rPr'))
        ea = None
        if rPr is not None:
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is not None:
                ea = rFonts.get(qn('w:eastAsia'))
        ea = ea or '(继承)'
        font_dist[ea] = font_dist.get(ea, 0) + 1
        if run.font.size:
            sz = run.font.size.pt
            size_dist[sz] = size_dist.get(sz, 0) + 1

print('eastAsia 字体分布:')
for fn, cnt in sorted(font_dist.items(), key=lambda x: -x[1]):
    print(f'  {fn:15s}: {cnt}')
print('\n字号 pt 分布:')
for sz, cnt in sorted(size_dist.items()):
    print(f'  {sz:>6.1f} pt: {cnt}')

# 正文中图/表引用核查
print('\n==== 正文中图/表引用核查 ====')
caption_indices = set()
for i, p in enumerate(paras):
    t = p.text.strip()
    if re.match(r'^[图表]\s*\d+', t):
        caption_indices.add(i)

fig_refs = {}
tbl_refs = {}
for i, p in enumerate(paras):
    if i in caption_indices:
        continue
    for m in re.finditer(r'图(\d+)', p.text):
        n = int(m.group(1))
        fig_refs[n] = fig_refs.get(n, 0) + 1
    for m in re.finditer(r'表(\d+)', p.text):
        n = int(m.group(1))
        tbl_refs[n] = tbl_refs.get(n, 0) + 1

print('图引用编号分布:')
for n in sorted(fig_refs.keys()):
    valid = '✓' if 1 <= n <= len(fig_nums) else '✗'
    print(f'  图{n}: {fig_refs[n]} 次 {valid}')

print('表引用编号分布:')
for n in sorted(tbl_refs.keys()):
    valid = '✓' if 1 <= n <= len(tbl_nums) else '✗'
    print(f'  表{n}: {tbl_refs[n]} 次 {valid}')
