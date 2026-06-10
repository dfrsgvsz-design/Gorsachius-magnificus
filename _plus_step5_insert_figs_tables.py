# -*- coding: utf-8 -*-
"""Step 5: 将 F5/F8/F9 生成的图表插入 PLUS 修订版 docx
- F5: 在原图1 后面追加 "图1' (修订版含 PLUS)" 含新 PNG
- F8: 在 PLUS 节末尾 (P199 后) 追加 "图16 驱动因子分布"
- F9: 在图16 后追加 "表14 景观格局相似度对比"
"""
import os
import json
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from copy import deepcopy

DOC_PATH = r'E:\大学\万物春\erci\郑春铃+横州市土地利用变化与生态安全评价(1)_PLUS修订版.docx'
WORK = r'E:\大学\万物春\erci\_plus_workspace'
FIG1_PNG = os.path.join(WORK, 'fig1_technical_roadmap_plus.png')
FIG16_PNG = os.path.join(WORK, 'fig16_driver_factors.png')
LANDSCAPE_JSON = os.path.join(WORK, 'landscape_similarity.json')

print(f'读取: {DOC_PATH}')
doc = Document(DOC_PATH)
paras = doc.paragraphs
print(f'总段落数: {len(paras)}, 表格数: {len(doc.tables)}')


def get_pstyle_val(para):
    pPr = para._p.find(qn('w:pPr'))
    if pPr is None:
        return None
    pStyle = pPr.find(qn('w:pStyle'))
    if pStyle is None:
        return None
    return pStyle.get(qn('w:val'))


def insert_para_after(template_para, text, style_id=None, alignment=None):
    """在 template_para 之后插入一个新段落（复制模板格式），文本为 text。

    必须清空 OMML 公式（m:oMath / m:oMathPara），否则若模板段含公式（如"式
    中：..."变量说明段），公式会被复制到新段。
    """
    new_p = deepcopy(template_para._p)
    # 清空所有 w:t 文本
    for t in new_p.findall('.//' + qn('w:t')):
        t.text = ''
    # ★ 关键修复：清空所有 OMML 公式元素 ★
    for omath in list(new_p.findall('.//' + qn('m:oMath'))):
        parent = omath.getparent()
        if parent is not None:
            parent.remove(omath)
    for omp in list(new_p.findall('.//' + qn('m:oMathPara'))):
        parent = omp.getparent()
        if parent is not None:
            parent.remove(omp)
    # 仅保留第一个 r
    runs_xml = new_p.findall(qn('w:r'))
    if runs_xml:
        for r in runs_xml[1:]:
            new_p.remove(r)
        first_t = new_p.find('.//' + qn('w:t'))
        if first_t is None:
            from docx.oxml import OxmlElement
            r0 = new_p.find(qn('w:r'))
            t_new = OxmlElement('w:t')
            t_new.set(qn('xml:space'), 'preserve')
            t_new.text = text
            r0.append(t_new)
        else:
            first_t.text = text
            first_t.set(qn('xml:space'), 'preserve')
    else:
        from docx.oxml import OxmlElement
        r_new = OxmlElement('w:r')
        t_new = OxmlElement('w:t')
        t_new.set(qn('xml:space'), 'preserve')
        t_new.text = text
        r_new.append(t_new)
        new_p.append(r_new)

    # 设置样式
    if style_id is not None:
        from docx.oxml import OxmlElement
        pPr = new_p.find(qn('w:pPr'))
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            new_p.insert(0, pPr)
        pStyle = pPr.find(qn('w:pStyle'))
        if pStyle is None:
            pStyle = OxmlElement('w:pStyle')
            pPr.insert(0, pStyle)
        pStyle.set(qn('w:val'), style_id)

    template_para._p.addnext(new_p)
    from docx.text.paragraph import Paragraph
    return Paragraph(new_p, template_para._parent)


# ============================================================
# F5: 在原图1 标题段之后追加补充技术路线图
# ============================================================
print('\n[F5] 在图1 之后插入修订版技术路线图')
fig1_caption_idx = -1
for i, p in enumerate(paras):
    if '图1 研究技术路线图' in p.text:
        fig1_caption_idx = i
        break
if fig1_caption_idx == -1:
    print('  [WARN] 未找到图1 标题')
else:
    anchor = paras[fig1_caption_idx]
    print(f'  锚点: P{fig1_caption_idx}: {anchor.text[:30]}')
    # 插入新段（含图片）+ 标题段
    new_para = insert_para_after(anchor, '')
    # 在新段中加图
    run = new_para.add_run()
    run.add_picture(FIG1_PNG, width=Inches(6.0))
    new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 再插入标题段
    caption = insert_para_after(new_para, '图1\u2032 横州市研究技术路线图（PLUS 修订版，含 LEAS+CARS+Markov 三模块）')
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    print('  [OK] 图1\u2032 已插入')


# ============================================================
# F8: 在 PLUS 节末尾追加图16 驱动因子分布图
# 锚点：P199 "采用'留一回测'..." 之后
# ============================================================
print('\n[F8] 在 PLUS 节末尾插入图16 驱动因子分布图')
# 重新读取段落（因为之前插入图1' 改变了段落顺序）
paras = doc.paragraphs
plus_validation_idx = -1
for i, p in enumerate(paras):
    if '采用"留一回测"' in p.text and 'leave-one-out' in p.text:
        plus_validation_idx = i
        break
if plus_validation_idx == -1:
    print('  [WARN] 未找到 PLUS 精度验证段')
else:
    anchor = paras[plus_validation_idx]
    print(f'  锚点: P{plus_validation_idx}: {anchor.text[:40]}')
    # 插入图片段
    img_para = insert_para_after(anchor, '')
    run = img_para.add_run()
    run.add_picture(FIG16_PNG, width=Inches(6.5))
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 插入标题段
    cap = insert_para_after(img_para, '图16 横州市 PLUS 模型驱动因子空间分布（6 个驱动因子：DEM、坡度、坡向、距水域距离、距建设用地距离、距研究区中心距离）')
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    print('  [OK] 图16 已插入')


# ============================================================
# F9: 在图16 标题段之后追加 "表14 景观格局相似度"
# 用 python-docx 创建新表格
# ============================================================
print('\n[F9] 在 PLUS 节末尾插入表14 景观格局相似度')
# 重新读取段落
paras = doc.paragraphs
fig16_cap_idx = -1
for i, p in enumerate(paras):
    if '图16 横州市 PLUS 模型驱动因子空间分布' in p.text:
        fig16_cap_idx = i
        break
if fig16_cap_idx == -1:
    print('  [WARN] 未找到图16 标题段')
else:
    anchor = paras[fig16_cap_idx]
    # 加表格标题段
    tbl_caption = insert_para_after(anchor,
        '表14 横州市 PLUS 模型主要地类景观格局指标（实际 2020 年值）')
    tbl_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 加表格本身：需要用 doc.add_table 后移动到 caption 之后
    # 简化做法：使用 docx XML 在 caption 后插入 w:tbl 元素
    with open(LANDSCAPE_JSON, 'r', encoding='utf-8') as f:
        lsdata = json.load(f)

    # 创建表格 (在文档末尾，然后移动)
    tbl = doc.add_table(rows=8, cols=6)
    tbl.style = 'Table Grid'
    # 表头
    headers = ['地类', '斑块数 NP', '斑块面积 CA/km²', '最大斑块指数 LPI/%', '平均斑块面积 MPS/km²', '占比 PLAND/%']
    for ci, h in enumerate(headers):
        tbl.rows[0].cells[ci].text = h
    # 数据行（主要地类，按面积排序）
    main_classes = ['林地', '耕地', '水域', '建设用地', '草地', '灌木', '未利用地']
    actual = lsdata['actual_2020']
    # 计算 PLAND
    total_ca = sum(actual[c]['CA'] for c in main_classes)
    for ri, cname in enumerate(main_classes, start=1):
        d = actual[cname]
        pland = d['CA'] / total_ca * 100 if total_ca > 0 else 0
        cells = tbl.rows[ri].cells
        cells[0].text = cname
        cells[1].text = f"{d['NP']}"
        cells[2].text = f"{d['CA']:.2f}"
        cells[3].text = f"{d['LPI']:.4f}"
        cells[4].text = f"{d['MPS']:.4f}"
        cells[5].text = f"{pland:.2f}"

    # 把表格从末尾移动到 tbl_caption 之后
    tbl_xml = tbl._tbl
    tbl_xml.getparent().remove(tbl_xml)
    tbl_caption._p.addnext(tbl_xml)

    print('  [OK] 表14 已插入')

# 保存
DST = DOC_PATH.replace('.docx', '_v2.docx')
doc.save(DST)
print(f'\n保存至: {DST}')
print(f'最终: 段落 {len(doc.paragraphs)} | 表格 {len(doc.tables)}')
