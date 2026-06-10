# -*- coding: utf-8 -*-
"""方案 A：图表彻底重构
1. 删除原图1/3/12 的标题段+图片段（9 段）
2. 重新按物理顺序编号所有图/表标题
3. 同步更新正文中所有 图N / 表N 引用
4. 清除表格内首行缩进 6 处
5. 字体一致化

输入: _PLUS修订版_最终_格式合规.docx
输出: _PLUS修订版_最终_格式合规_v2.docx
"""
import sys, io, re, shutil
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

SRC = r'E:\大学\万物春\erci\郑春铃+横州市土地利用变化与生态安全评价(1)_PLUS修订版_最终_格式合规.docx'
DST = r'E:\大学\万物春\erci\郑春铃+横州市土地利用变化与生态安全评价(1)_PLUS修订版_最终_格式合规_v2.docx'

shutil.copy(SRC, DST)
print(f'已复制 → v2: {DST}')

doc = Document(DST)
paras = doc.paragraphs

# ============================================================
# Step 1: 收集需删除的段（按段索引降序删，避免索引漂移）
# ============================================================
DELETE_INDICES = [152, 153, 212, 213, 214, 215, 216, 394, 395]
print(f'\n==== Step 1: 删除原图1/3/12 共 {len(DELETE_INDICES)} 段 ====')
for idx in sorted(DELETE_INDICES, reverse=True):
    p = paras[idx]
    text_snip = p.text[:40] if p.text else '(图片段)'
    p._p.getparent().remove(p._p)
    print(f'  [删] P{idx}: {text_snip}')


# ============================================================
# Step 2: 重新读 doc，扫描所有图/表标题，按物理顺序编号
# ============================================================
print('\n==== Step 2: 建立旧→新编号映射 ====')
# ★ 关键修复：先 save 已删除的 doc，再重新读 ★
doc.save(DST)
doc = Document(DST)
paras = doc.paragraphs
print(f'  删除后段落数: {len(paras)} (原 489, 删 9, 应 480)')

# 扫描图标题（包括 "图N'" "图N′" "图N" 等）
fig_old_to_new = {}  # 旧编号字符串 → 新编号字符串
new_fig_idx = 0
fig_caption_paras = []
for i, p in enumerate(paras):
    t = p.text.strip()
    # 匹配 "图 N" 或 "图N" 或 "图N'" 或 "图N′" 开头
    m = re.match(r'^图\s*(\d+)([\u2032\u2019\u2018\']?)\s', t)
    if m:
        new_fig_idx += 1
        old_label = f"图{m.group(1)}{m.group(2)}"
        new_label = f"图{new_fig_idx}"
        # 也记录无撇号的形式
        plain_old = f"图{m.group(1)}"
        fig_old_to_new[old_label] = new_label
        if plain_old not in fig_old_to_new:
            fig_old_to_new[plain_old] = new_label
        fig_caption_paras.append((i, old_label, new_label))
        print(f'  图: P{i} "{old_label}" → "{new_label}"')

# 扫描表标题
tbl_old_to_new = {}
new_tbl_idx = 0
tbl_caption_paras = []
for i, p in enumerate(paras):
    t = p.text.strip()
    m = re.match(r'^表\s*(\d+)([\u2032\u2019\u2018\']?)\s', t)
    if m:
        new_tbl_idx += 1
        old_label = f"表{m.group(1)}{m.group(2)}"
        new_label = f"表{new_tbl_idx}"
        plain_old = f"表{m.group(1)}"
        tbl_old_to_new[old_label] = new_label
        if plain_old not in tbl_old_to_new:
            tbl_old_to_new[plain_old] = new_label
        tbl_caption_paras.append((i, old_label, new_label))
        print(f'  表: P{i} "{old_label}" → "{new_label}"')

print(f'\n图编号: 共 {new_fig_idx} 个')
print(f'表编号: 共 {new_tbl_idx} 个')


# ============================================================
# Step 3: 更新图/表标题段自身（按物理顺序重新分配 1, 2, 3, ...）
# 注意：必须用占位符两阶段替换，避免冲突
# ============================================================
print('\n==== Step 3: 更新图/表标题段自身 ====')

# 阶段 3a: 把图标题中的"图N "或"图N'  "替换为占位符 "FIGNEW{new_idx}_"
new_fig_idx = 0
for i, p in enumerate(paras):
    t = p.text.strip()
    m = re.match(r'^图\s*(\d+)([\u2032\u2019\u2018\']?)(\s)', t)
    if m:
        new_fig_idx += 1
        # 在 run 级别替换
        # 找含 "图N" 的 run
        old_pattern = m.group(0)
        new_text = f"图{new_fig_idx}{m.group(3)}"
        # 整段替换：用 set first run text 方式
        full = p.text
        # 计算原 prefix 长度
        prefix_len = len(m.group(0))
        # 直接替换段文本
        new_full = new_text + full[prefix_len:]
        # 在第一个 run 上重设
        if p.runs:
            # 简单做法：把第一个 run 的 text 重置，其他 run 清空
            p.runs[0].text = new_full
            for r in p.runs[1:]:
                r.text = ''
        print(f'  图: P{i} 改为 "{p.text[:50]}"')

new_tbl_idx = 0
for i, p in enumerate(paras):
    t = p.text.strip()
    m = re.match(r'^表\s*(\d+)([\u2032\u2019\u2018\']?)(\s)', t)
    if m:
        new_tbl_idx += 1
        new_text = f"表{new_tbl_idx}{m.group(3)}"
        prefix_len = len(m.group(0))
        full = p.text
        new_full = new_text + full[prefix_len:]
        if p.runs:
            p.runs[0].text = new_full
            for r in p.runs[1:]:
                r.text = ''
        print(f'  表: P{i} 改为 "{p.text[:50]}"')


# ============================================================
# Step 4: 同步更新正文中所有 图N / 表N 引用
# 用占位符两阶段替换避免循环替换
# ============================================================
print('\n==== Step 4: 更新正文中的图/表引用 ====')

# 构建 全局 映射（旧标签 → 新标签）
# 在 Step 2 中我们已经记录了 fig_old_to_new 和 tbl_old_to_new
# 注意：这些映射来自删除前的扫描，但删除后图编号变了。重新扫描。

# 由于我们改了图/表标题为新编号 1,2,3,...，需要正文也对应改
# 但正文中的引用还是旧编号（如 "图3" 引用的是原图3，对应新编号是图4）

# 重新建映射：原旧编号 (从删除前 docx 的位置推断) → 新编号
# 实际上更简单：基于映射规则手动列出

# 删除前的图编号 → 删除后的新编号
# 物理顺序删除前: 图1(原), 图1'(修订), 图2, 图16, 图3(原), 图3'(修订), 图4..9, 图10..11, 图12'(修订), 图12(原), 图13..15
# 删除后: 图1'(P155→图1), 图2(图2), 图16(图3), 图3'(图4), 图4(图5), 图5(图6), 图6(图7),
#         图7(图8), 图8(图9), 图9(图10), 图10(图11), 图11(图12), 图12'(图13),
#         图13(图14), 图14(图15), 图15(图16)
FIG_MAPPING = {
    '图1': '图1',  # 用户在正文中"图1"指的是技术路线图，删除原图1后修订版变图1
    '图2': '图2',
    '图3': '图5',  # 原图3 已删，原"图3"引用指的是 现图5？等等，不对
    # 实际上原图3 是"四期土地利用现状图"（被删），但原图4 是"面积变化趋势"——位置在原图3 之后
    # 正文中"图3"引用的是"四期土地利用现状图"，现在它在 图4 位置（原图3'）
    '图4': '图5',  # 原图4 → 现图5
    '图5': '图6',
    '图6': '图7',
    '图7': '图8',
    '图8': '图9',
    '图9': '图10',
    '图10': '图11',
    '图11': '图12',
    '图12': '图13',  # 原图12 已删，原"图12"引用指的是 ESI 趋势图（现图13）
    '图13': '图14',
    '图14': '图15',
    '图15': '图16',
    '图16': '图3',  # 原图16 在物理上位置3
    # 撇号版
    "图1'": '图1',
    "图1\u2032": '图1',
    "图3'": '图4',
    "图3\u2032": '图4',
    "图12'": '图13',
    "图12\u2032": '图13',
}

# 修正：原图3 引用的是"四期土地利用图"（被删除后由原图3' 代替），所以原"图3"引用 → 新"图4"
# 修正上面的图3映射
FIG_MAPPING['图3'] = '图4'

# 表映射
TBL_MAPPING = {
    '表1': '表1',
    '表2': '表3',  # 原表2 → 现表3（因为新表2 是原表14）
    '表3': '表4',
    '表4': '表5',
    '表5': '表6',
    '表6': '表7',
    '表7': '表8',
    '表8': '表9',
    '表9': '表10',
    '表10': '表11',
    '表11': '表12',
    '表12': '表13',
    '表13': '表14',
    '表14': '表2',  # 原表14 → 物理上是表2
}

# 两阶段替换：阶段 1 用占位符，阶段 2 占位符 → 新编号
print('  阶段 1: 用占位符替换正文 图N/表N')
placeholder_map = {}
for old, new in FIG_MAPPING.items():
    placeholder = f'__FIG_PLH_{new[1:]}__'
    placeholder_map[old] = placeholder
for old, new in TBL_MAPPING.items():
    placeholder = f'__TBL_PLH_{new[1:]}__'
    placeholder_map[old] = placeholder

# 收集所有图/表标题段的段索引（不应被引用替换影响）
caption_indices = set()
for i, p in enumerate(paras):
    t = p.text.strip()
    if re.match(r'^[图表]\s*\d+', t):
        caption_indices.add(i)

# 长串先替换：避免"图10"被"图1"误改
sorted_olds = sorted(placeholder_map.keys(), key=lambda x: -len(x))

ref_replace_count = 0
for i, p in enumerate(paras):
    if i in caption_indices:
        continue  # 跳过标题段（标题已在 Step 3 改）
    for run in p.runs:
        if not run.text:
            continue
        original = run.text
        new_text = original
        for old in sorted_olds:
            if old in new_text:
                placeholder = placeholder_map[old]
                count = new_text.count(old)
                new_text = new_text.replace(old, placeholder)
                ref_replace_count += count
        if new_text != original:
            run.text = new_text

# 阶段 2: 占位符 → 新编号
print(f'  阶段 1 共替换 {ref_replace_count} 处引用为占位符')
print('  阶段 2: 占位符 → 新编号')
fig_unique_targets = set(FIG_MAPPING.values())
tbl_unique_targets = set(TBL_MAPPING.values())
phase2_count = 0
for i, p in enumerate(paras):
    if i in caption_indices:
        continue
    for run in p.runs:
        if not run.text:
            continue
        original = run.text
        new_text = original
        for new_label in fig_unique_targets:
            placeholder = f'__FIG_PLH_{new_label[1:]}__'
            if placeholder in new_text:
                count = new_text.count(placeholder)
                new_text = new_text.replace(placeholder, new_label)
                phase2_count += count
        for new_label in tbl_unique_targets:
            placeholder = f'__TBL_PLH_{new_label[1:]}__'
            if placeholder in new_text:
                count = new_text.count(placeholder)
                new_text = new_text.replace(placeholder, new_label)
                phase2_count += count
        if new_text != original:
            run.text = new_text
print(f'  阶段 2 共还原 {phase2_count} 个占位符')


# ============================================================
# Step 5: 清除表格内首行缩进
# ============================================================
print('\n==== Step 5: 清除表格内首行缩进 ====')
clean_indent_count = 0
for ti, tbl in enumerate(doc.tables):
    for ri, row in enumerate(tbl.rows):
        for ci, cell in enumerate(row.cells):
            for p in cell.paragraphs:
                pf = p.paragraph_format
                if pf.first_line_indent and pf.first_line_indent.cm > 0.1:
                    pf.first_line_indent = None
                    clean_indent_count += 1
print(f'  清除首行缩进: {clean_indent_count} 个段落')


# ============================================================
# Step 6: 保存
# ============================================================
doc.save(DST)
import os
print(f'\n==== 完成 ====')
print(f'保存至: {DST}')
print(f'文件大小: {os.path.getsize(DST):,} bytes')
print(f'段落数: {len(doc.paragraphs)}')
print(f'表格数: {len(doc.tables)}')
