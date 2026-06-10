# -*- coding: utf-8 -*-
"""综合修复 5 项格式违规
F1: 参考文献补 [N] 编号 (P453-P480)
F2: 一级标题字号 → 小二号(18pt) 黑体
F3: 正文 Normal 段字号字体 → 小四号(12pt) 宋体（仅 P103 之后绪论章起）
F4: 16 个表格全部改为三线表
F5: 中文段中英文逗号 → 中文逗号

输入: _PLUS修订版_最终.docx
输出: _PLUS修订版_最终_格式合规.docx
"""
import sys, io, re, shutil
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = r'E:\大学\万物春\erci\郑春铃+横州市土地利用变化与生态安全评价(1)_PLUS修订版_最终.docx'
DST = r'E:\大学\万物春\erci\郑春铃+横州市土地利用变化与生态安全评价(1)_PLUS修订版_最终_格式合规.docx'

shutil.copy(SRC, DST)
print(f'已复制 → {DST}')
doc = Document(DST)
paras = doc.paragraphs


def set_run_font_size_chinese(run, pt_size, cn_font, west_font='Times New Roman'):
    """设置 run 的字号 + 中英文字体（标准做法）。"""
    run.font.size = Pt(pt_size)
    run.font.name = west_font  # 西文
    # 设置中文字体 eastAsia
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), cn_font)
    rFonts.set(qn('w:ascii'), west_font)
    rFonts.set(qn('w:hAnsi'), west_font)


# ============================================================
# F1: 参考文献补 [N] 编号
# ============================================================
print('\n==== F1: 参考文献补 [N] 编号 ====')
ref_start = -1
for i, p in enumerate(paras):
    if p.text.strip() == '参考文献':
        ref_start = i
        break

if ref_start < 0:
    print('  [WARN] 未找到参考文献节')
else:
    print(f'  参考文献节起点: P{ref_start}')
    ref_num = 0
    for i in range(ref_start + 1, len(paras)):
        p = paras[i]
        t = p.text.strip()
        if not t:
            continue
        if t in ('致谢', '附录'):
            break
        ref_num += 1
        # 检查是否已有编号
        if re.match(r'^\[\d+\]', t):
            print(f'  [SKIP] P{i} 已有编号: {t[:50]}')
            continue
        # 在段首插入 [N] 编号
        # 找第一个 run，前插一个新 run 包含 "[N] "
        prefix = f'[{ref_num}] '
        # 直接修改第一个 run 的 text
        if p.runs:
            first_run = p.runs[0]
            # 复制其格式新建 run，但更简单做法：直接前插字符
            first_run.text = prefix + first_run.text
        else:
            # 无 run，创建一个
            r_new = OxmlElement('w:r')
            t_new = OxmlElement('w:t')
            t_new.set(qn('xml:space'), 'preserve')
            t_new.text = prefix + t
            r_new.append(t_new)
            p._p.append(r_new)
        print(f'  [F1] P{i} → [{ref_num}] {t[:60]}')

    print(f'  总计补编号: {ref_num} 条')


# ============================================================
# F2: 一级标题字号 → 小二号 18pt 黑体居中
# ============================================================
print('\n==== F2: 一级标题字号 ====')
h1_count = 0
for p in doc.paragraphs:
    if p.style.name == 'Heading 1' and p.text.strip():
        h1_count += 1
        for run in p.runs:
            set_run_font_size_chinese(run, 18, '黑体')
        # 居中对齐
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 行距设置（如果有需要可加）
print(f'  修改 Heading 1 段: {h1_count} 个')


# ============================================================
# F3: 正文 Normal 段字号字体 → 小四号 12pt 宋体
# 策略：仅修改 P103（绪论起点）及之后的 Normal 段，避免动诚信承诺书
# ============================================================
print('\n==== F3: 正文 Normal 字号字体 ====')
intro_start = -1
for i, p in enumerate(doc.paragraphs):
    if '一、绪论' in p.text or p.text.strip() == '一、绪论':
        intro_start = i
        break
print(f'  绪论起点: P{intro_start}')

if intro_start < 0:
    print('  [WARN] 未找到绪论起点，跳过 F3')
else:
    normal_fixed = 0
    paras2 = doc.paragraphs
    for i in range(intro_start, len(paras2)):
        p = paras2[i]
        # 跳过参考文献节及之后（参考文献用小四号仿宋）
        if p.text.strip() == '参考文献':
            print(f'  P{i} 到达参考文献节，正文修改终止')
            break
        # 仅修改 Normal 样式段
        if p.style.name == 'Normal' and p.runs:
            for run in p.runs:
                # 仅当该 run 有文本时才设
                if run.text.strip():
                    set_run_font_size_chinese(run, 12, '宋体')
            normal_fixed += 1
    print(f'  修改 Normal 正文段: {normal_fixed} 个')


# ============================================================
# F4: 表格转为三线表
# ============================================================
print('\n==== F4: 16 个表格转为三线表 ====')


def set_cell_borders_three_line(cell, position, is_first_row=False, is_last_row=False,
                                 is_header_row=False):
    """三线表：仅顶部、表头下方、底部有线，其他无线。"""
    tcPr = cell._tc.get_or_add_tcPr()
    # 先移除现有 tcBorders
    existing = tcPr.find(qn('w:tcBorders'))
    if existing is not None:
        tcPr.remove(existing)
    tcBorders = OxmlElement('w:tcBorders')

    # 顶部边框
    top = OxmlElement('w:top')
    if is_first_row:
        top.set(qn('w:val'), 'single')
        top.set(qn('w:sz'), '18')  # 1.5pt（粗）
        top.set(qn('w:color'), '000000')
    else:
        top.set(qn('w:val'), 'nil')
    tcBorders.append(top)

    # 底部边框
    bottom = OxmlElement('w:bottom')
    if is_last_row:
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '18')  # 1.5pt（粗）
        bottom.set(qn('w:color'), '000000')
    elif is_header_row:
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '8')  # 0.5pt（细）
        bottom.set(qn('w:color'), '000000')
    else:
        bottom.set(qn('w:val'), 'nil')
    tcBorders.append(bottom)

    # 左右无线
    for side in ('left', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'nil')
        tcBorders.append(el)

    tcPr.append(tcBorders)


for ti, tbl in enumerate(doc.tables):
    n_rows = len(tbl.rows)
    if n_rows == 0:
        continue
    # 移除表格级 style（避免 Grid 样式覆盖）
    tbl.style = doc.styles['Normal Table'] if 'Normal Table' in [s.name for s in doc.styles] else tbl.style
    # 清除表格级边框
    tblPr = tbl._tbl.find(qn('w:tblPr'))
    if tblPr is not None:
        existing_borders = tblPr.find(qn('w:tblBorders'))
        if existing_borders is not None:
            tblPr.remove(existing_borders)
        # 添加 tblBorders 标记无线
        tblBorders = OxmlElement('w:tblBorders')
        for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), 'nil')
            tblBorders.append(el)
        tblPr.append(tblBorders)

    # 为每个单元格设三线表边框
    for ri, row in enumerate(tbl.rows):
        is_first = (ri == 0)
        is_last = (ri == n_rows - 1)
        is_header = (ri == 0)  # 表头通常是第 0 行
        for cell in row.cells:
            set_cell_borders_three_line(cell, ri, is_first_row=is_first,
                                         is_last_row=is_last,
                                         is_header_row=is_header)
print(f'  转为三线表: {len(doc.tables)} 个')


# ============================================================
# F5: 中文段中英文逗号 → 中文逗号
# ============================================================
print('\n==== F5: 中文段英文逗号 → 中文逗号 ====')
replaced_count = 0
for p in doc.paragraphs:
    t = p.text
    if not t.strip():
        continue
    cn_chars = len(re.findall(r'[\u4e00-\u9fff]', t))
    if cn_chars < len(t) * 0.5:
        continue  # 非中文主体段（如英文摘要、参考文献英文条目），跳过
    # 在 run 级别替换
    for run in p.runs:
        if not run.text:
            continue
        new_text = re.sub(r',([ \u4e00-\u9fff])', r'，\1', run.text)
        if new_text != run.text:
            # 计算替换数
            diff = sum(1 for a, b in zip(run.text, new_text) if a != b)
            replaced_count += diff
            run.text = new_text
print(f'  替换中文段英文逗号: {replaced_count} 处')


# ============================================================
# 保存
# ============================================================
doc.save(DST)
import os
print(f'\n保存至: {DST}')
print(f'文件大小: {os.path.getsize(DST):,} bytes')
print(f'段落数: {len(doc.paragraphs)}')
print(f'表格数: {len(doc.tables)}')
