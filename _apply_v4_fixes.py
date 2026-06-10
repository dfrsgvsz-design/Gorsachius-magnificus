# -*- coding: utf-8 -*-
"""
v3 → v4 综合修复：
F-DUP   移除 28 个参考文献段的 numPr 属性（消除 [N][N] 双重编号）
F-FIG13 重画图13 ESI 折线图（去掉内嵌"图12"标题），替换 docx 中嵌入的图片
"""
import os, sys, io, re, zipfile, shutil
sys.stdout = io.open(r'f:\Gorsachius magnificus\_apply_v4_log.txt','w',encoding='utf-8')

import numpy as np
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from matplotlib import rcParams

rcParams['font.sans-serif'] = ['Microsoft YaHei', 'SimHei', 'DejaVu Sans']
rcParams['axes.unicode_minus'] = False
rcParams['font.size'] = 10

WORK = r'E:\大学\万物春\erci\_plus_workspace'
SRC = r'E:\大学\万物春\erci\郑春铃+横州市土地利用变化与生态安全评价(1)_PLUS修订版_最终_格式合规_v3.docx'
DST = r'E:\大学\万物春\erci\郑春铃+横州市土地利用变化与生态安全评价(1)_PLUS修订版_最终_格式合规_v4.docx'

# ============================================================
# F-FIG13: 重画 ESI 折线图（无内嵌"图12"标题）
# ============================================================
def make_fig13_clean():
    """重画 ESI 演变图，移除内嵌的'图12 ...'标题。
    标题由 docx 外的图标题段承担。"""
    years = [2010, 2015, 2020, 2025]
    esi = [0.718, 0.460, 0.390, 0.283]
    levels = ['Ⅱ 较安全', 'Ⅲ 预警', 'Ⅳ 中度不安全', 'Ⅳ 中度不安全']

    fig, ax = plt.subplots(figsize=(10, 6), dpi=200)

    ax.plot(years[:3], esi[:3], 'o-', color='#1E40AF', linewidth=2.5,
            markersize=12, markerfacecolor='#3B82F6', markeredgecolor='#1E40AF',
            markeredgewidth=2, label='实测情景 (2010—2020)', zorder=3)
    ax.plot(years[2:], esi[2:], '--', color='#F59E0B', linewidth=2.5, zorder=2)
    ax.plot(2025, 0.283, 'o', markersize=14, markerfacecolor='white',
            markeredgecolor='#D97706', markeredgewidth=3,
            label='2025 PLUS 模拟情景', zorder=4)

    for x, y, lvl in zip(years, esi, levels):
        ax.annotate(f'ESI={y:.3f}\n({lvl})',
                    xy=(x, y), xytext=(0, 18), textcoords='offset points',
                    ha='center', fontsize=10,
                    bbox=dict(boxstyle='round,pad=0.3',
                              facecolor='white' if x != 2025 else '#FEF3C7',
                              edgecolor='#D97706' if x == 2025 else '#94A3B8',
                              linewidth=1))

    level_thresholds = {0.8: 'Ⅰ 安全', 0.6: 'Ⅱ 较安全', 0.4: 'Ⅲ 预警', 0.2: 'Ⅳ 中度不安全'}
    for thr, name in level_thresholds.items():
        ax.axhline(y=thr, color='gray', linestyle=':', linewidth=0.8, alpha=0.5)
        ax.text(2009.5, thr + 0.01, name, fontsize=8, color='gray', alpha=0.7)

    ax.axvspan(2020, 2025, alpha=0.08, color='#F59E0B', label='PLUS 模拟期')

    ax.set_xlabel('年份', fontsize=12, fontweight='bold')
    ax.set_ylabel('生态安全指数 ESI', fontsize=12, fontweight='bold')
    # 关键：移除 set_title 调用，不绘制内嵌标题
    ax.set_xticks(years)
    ax.set_xlim(2008, 2027)
    ax.set_ylim(0.15, 0.85)
    ax.grid(True, alpha=0.3)
    ax.legend(loc='upper right', fontsize=10, framealpha=0.95)

    plt.tight_layout()
    out_path = os.path.join(WORK, 'fig13_esi_clean.png')
    plt.savefig(out_path, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close()
    print(f"[F-FIG13] 新图: {out_path} ({os.path.getsize(out_path)} bytes)")
    return out_path


# ============================================================
# 工具函数：在 docx zip 内列出/读取/替换 media 图片
# ============================================================
def list_docx_images(docx_path):
    """列出 docx zip 内所有 word/media/* 图片。"""
    images = []
    with zipfile.ZipFile(docx_path, 'r') as z:
        for name in z.namelist():
            if name.startswith('word/media/'):
                images.append((name, z.getinfo(name).file_size))
    return images


def replace_docx_image(docx_in, docx_out, target_name, new_png_path):
    """把 docx_in 中 word/media/<target_name> 替换为 new_png_path，写到 docx_out。"""
    with open(new_png_path, 'rb') as f:
        new_data = f.read()
    with zipfile.ZipFile(docx_in, 'r') as zin:
        with zipfile.ZipFile(docx_out, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                if item.filename == target_name:
                    zout.writestr(item, new_data)
                else:
                    zout.writestr(item, zin.read(item.filename))


def find_fig13_image_name(docx_path):
    """通过分析 docx XML，找出图13 标题段附近段落引用的图片资源 ID，
    然后查找该 ID 对应的 media 文件名。"""
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document(docx_path)
    paras = doc.paragraphs

    # 找图13 标题段
    fig13_idx = -1
    for i, p in enumerate(paras):
        t = p.text.strip()
        if re.match(r'^图\s*13[\s：:]', t):
            fig13_idx = i
            break
    if fig13_idx < 0:
        return None, []
    print(f"[F-FIG13] 图13 标题段位于 P{fig13_idx}: {paras[fig13_idx].text[:80]!r}")

    # 在 P(fig13_idx-1) 或 P(fig13_idx-2) 中找 blip 引用
    NS_R = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
    NS_A = 'http://schemas.openxmlformats.org/drawingml/2006/main'
    candidates = []
    for offset in [-3, -2, -1, 1, 2]:
        pi = fig13_idx + offset
        if pi < 0 or pi >= len(paras):
            continue
        el = paras[pi]._element
        blips = el.findall(f'.//{{{NS_A}}}blip')
        for blip in blips:
            rid = blip.get(f'{{{NS_R}}}embed')
            if rid:
                candidates.append((pi, rid))
    print(f"[F-FIG13] 候选 blip rId: {candidates}")
    if not candidates:
        return None, []

    # 通过 relationships 找 image 文件名
    # 加载 word/_rels/document.xml.rels
    rid_to_target = {}
    with zipfile.ZipFile(docx_path, 'r') as z:
        with z.open('word/_rels/document.xml.rels') as f:
            rels_xml = f.read().decode('utf-8')
    from xml.etree import ElementTree as ET
    ns_pkg = 'http://schemas.openxmlformats.org/package/2006/relationships'
    root = ET.fromstring(rels_xml)
    for rel in root.findall(f'{{{ns_pkg}}}Relationship'):
        rid_to_target[rel.get('Id')] = rel.get('Target')

    chosen_pi, chosen_rid = candidates[0]
    target = rid_to_target.get(chosen_rid)
    if target and not target.startswith('word/'):
        target = 'word/' + target
    print(f"[F-FIG13] 选定 rId={chosen_rid}, target={target}")
    return target, candidates


# ============================================================
# F-DUP: 用 python-docx 移除 28 个文献段的 numPr
# ============================================================
def remove_ref_numbering(docx_path, out_path):
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document(docx_path)
    paras = doc.paragraphs
    removed = 0
    for pi, p in enumerate(paras):
        t = p.text.strip()
        if not re.match(r'^\[(\d+)\]', t):
            continue
        # 跳过 'toc N' 样式段
        if 'toc' in p.style.name.lower():
            continue
        pPr = p._element.find(qn('w:pPr'))
        if pPr is None:
            continue
        numPr = pPr.find(qn('w:numPr'))
        if numPr is not None:
            pPr.remove(numPr)
            removed += 1
    print(f"[F-DUP] 移除 {removed} 个文献段的 numPr")
    doc.save(out_path)
    print(f"[F-DUP] 保存至 {out_path}")
    return removed


# ============================================================
# 主流程
# ============================================================
def main():
    print(f"开始: v3 -> v4")
    print(f"SRC: {SRC}")
    print(f"DST: {DST}")
    print()

    # Step 1: 重画 fig13 png
    print("=" * 60)
    print("Step 1: 重画图13 (无内嵌标题)")
    print("=" * 60)
    new_png = make_fig13_clean()

    # Step 2: 找出 docx 内图13 对应的 media 文件名
    print("\n" + "=" * 60)
    print("Step 2: 定位图13 在 docx 内的图片资源")
    print("=" * 60)
    target_name, candidates = find_fig13_image_name(SRC)
    if not target_name:
        print("  ✗ 未定位到图13 的图片资源，跳过 F-FIG13")
    else:
        # 列出 docx 内全部 media
        all_imgs = list_docx_images(SRC)
        print(f"  docx 内 media 文件 ({len(all_imgs)}):")
        for n, sz in all_imgs:
            mark = " ← 目标" if n == target_name else ""
            print(f"    {n}  {sz} bytes{mark}")

    # Step 3: 先做 F-DUP（python-docx 操作），再做 F-FIG13（zip 操作）
    print("\n" + "=" * 60)
    print("Step 3: F-DUP 移除 28 个文献段 numPr")
    print("=" * 60)
    tmp_path = DST.replace('.docx', '_tmp.docx')
    n_removed = remove_ref_numbering(SRC, tmp_path)

    # Step 4: F-FIG13 替换图片
    print("\n" + "=" * 60)
    print("Step 4: F-FIG13 替换 docx 内的图13 图片")
    print("=" * 60)
    if target_name:
        replace_docx_image(tmp_path, DST, target_name, new_png)
        os.remove(tmp_path)
        print(f"  [OK] 已替换 {target_name}")
    else:
        shutil.move(tmp_path, DST)

    print(f"\n最终输出: {DST}")
    print(f"文件大小: {os.path.getsize(DST):,} bytes")
    print("\nDone.")


if __name__ == '__main__':
    main()
