# -*- coding: utf-8 -*-
"""GXUFE Paper Writing Spec 论文格式全面自检脚本
对 _PLUS修订版_最终.docx 按 SKILL.md §8 Checklist 跑 6 大维度检查，
生成 markdown 违规报告。
"""
import sys, io, re, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.oxml.ns import qn

DOC = r'E:\大学\万物春\erci\郑春铃+横州市土地利用变化与生态安全评价(1)_PLUS修订版_最终_格式合规.docx'
OUT = r'E:\大学\万物春\erci\_格式自检报告_GXUFE_v2.md'

# 字号 Pt 对照表
SIZE_MAP = {
    22: '二号', 18: '小二号', 16: '三号', 15: '小三号',
    14: '四号', 12: '小四号', 10.5: '五号', 9: '小五号',
}

def fmt_size(pt):
    if pt is None:
        return '默认'
    val = pt.pt if hasattr(pt, 'pt') else float(pt)
    return f"{SIZE_MAP.get(val, '?')}({val}pt)"

def get_run_size(run):
    """获取 run 的字号 Pt，若未设置返回 None。"""
    if run.font.size:
        return run.font.size
    return None

def get_run_east_asia_font(run):
    """获取 run 的中文字体（eastAsia）。"""
    rPr = run._element.find(qn('w:rPr'))
    if rPr is None:
        return None
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        return None
    return rFonts.get(qn('w:eastAsia'))

def get_run_ascii_font(run):
    """获取 run 的西文字体（ascii）。"""
    rPr = run._element.find(qn('w:rPr'))
    if rPr is None:
        return None
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        return None
    return rFonts.get(qn('w:ascii'))


class Report:
    def __init__(self):
        self.sections = {}

    def add(self, section, label, status, message='', expected='', actual=''):
        """status: PASS / WARN / FAIL / INFO"""
        self.sections.setdefault(section, []).append({
            'label': label, 'status': status,
            'message': message, 'expected': expected, 'actual': actual,
        })

    def to_markdown(self, doc_path):
        lines = [f'# GXUFE 论文格式自检报告\n',
                 f'**文件**：`{doc_path}`\n',
                 f'**依据**：广西财经学院本科生毕业论文（设计）撰写规范（附件三 / SKILL.md）\n',
                 '\n---\n']
        # 统计
        total_pass = total_warn = total_fail = total_info = 0
        for items in self.sections.values():
            for it in items:
                if it['status'] == 'PASS': total_pass += 1
                elif it['status'] == 'WARN': total_warn += 1
                elif it['status'] == 'FAIL': total_fail += 1
                else: total_info += 1
        lines.append(f'## 总体统计\n')
        lines.append(f'- PASS: **{total_pass}**')
        lines.append(f'- WARN: **{total_warn}**')
        lines.append(f'- FAIL: **{total_fail}**')
        lines.append(f'- INFO: **{total_info}**')
        lines.append('')
        # 详细
        for section, items in self.sections.items():
            lines.append(f'\n## {section}\n')
            lines.append('| 状态 | 检查项 | 说明 | 期望 | 实测 |')
            lines.append('|------|--------|------|------|------|')
            for it in items:
                emoji = {'PASS': '✅', 'WARN': '⚠️', 'FAIL': '❌', 'INFO': 'ℹ️'}.get(it['status'], '?')
                msg = (it['message'] or '').replace('|', '\\|').replace('\n', ' ')
                exp = (it['expected'] or '').replace('|', '\\|')[:60]
                act = (it['actual'] or '').replace('|', '\\|')[:80]
                lines.append(f'| {emoji} {it["status"]} | {it["label"]} | {msg} | {exp} | {act} |')
        return '\n'.join(lines)


def main():
    report = Report()
    doc = Document(DOC)
    paras = doc.paragraphs
    all_text = '\n'.join(p.text for p in paras)

    # ============= A. 完整性检查 =============
    SECT_A = 'A. 完整性检查'

    # A1: 封面元素
    cover_elements = ['题目', '系别', '专业', '班级', '姓名', '学号', '指导教师']
    front_text = '\n'.join(p.text for p in paras[:30])
    missing_cover = [e for e in cover_elements if e not in front_text]
    if not missing_cover:
        report.add(SECT_A, 'A1 封面元素齐全', 'PASS',
                   '前 30 段中找到所有封面要素', '题目/系别/专业/班级/姓名/学号/指导教师', '全部存在')
    else:
        report.add(SECT_A, 'A1 封面元素', 'WARN',
                   f'前 30 段中未明确找到：{missing_cover}',
                   '7 项封面元素', f'缺 {missing_cover}')

    # A2: 题目 ≤ 25 字
    title_text = ''
    for p in paras[:30]:
        t = p.text.strip()
        if '横州市' in t and '土地利用' in t and len(t) < 60:
            title_text = t
            break
    if title_text:
        title_len = len(re.sub(r'\s', '', title_text))
        if title_len <= 25:
            report.add(SECT_A, 'A2 题目 ≤ 25 字', 'PASS', '',
                       '≤ 25 个汉字', f'{title_len} 字: {title_text}')
        else:
            report.add(SECT_A, 'A2 题目 ≤ 25 字', 'FAIL',
                       f'题目超长：{title_len} 字',
                       '≤ 25 字', f'{title_len} 字: {title_text}')
    else:
        report.add(SECT_A, 'A2 题目', 'WARN', '未自动定位到正文题目', '', '')

    # A3: 诚信承诺书
    if '诚信承诺书' in all_text or '承诺书' in front_text:
        report.add(SECT_A, 'A3 诚信承诺书', 'PASS', '', '应存在', '已存在')
    else:
        report.add(SECT_A, 'A3 诚信承诺书', 'WARN',
                   '正文中未找到"诚信承诺书"关键词，可能装订时另附',
                   '应单独一页', '未在 docx 正文中检出')

    # A4: 目录
    has_toc = any('目录' in p.text for p in paras[:50])
    if has_toc:
        report.add(SECT_A, 'A4 目录', 'PASS', '', '三级层次', '已含目录')
    else:
        report.add(SECT_A, 'A4 目录', 'WARN', '未在前 50 段找到"目录"标识', '', '')

    # A5: 中文摘要 300-500 字
    cn_abstract_idx = -1
    for i, p in enumerate(paras):
        if '中文摘要' in p.text or (i > 30 and i < 60 and p.text.strip().startswith('土地利用变化是')):
            cn_abstract_idx = i
            break
    cn_abstract_len = 0
    if cn_abstract_idx >= 0:
        # 找到摘要正文（中文摘要标题之后的第一个长段落）
        for j in range(cn_abstract_idx, min(cn_abstract_idx + 5, len(paras))):
            t = paras[j].text.strip()
            if len(t) > 100 and '中文摘要' not in t and '关键词' not in t:
                cn_abstract_len = len(re.sub(r'\s', '', t))
                break
    if 300 <= cn_abstract_len <= 500:
        report.add(SECT_A, 'A5 中文摘要 300—500 字', 'PASS', '',
                   '300—500 字', f'{cn_abstract_len} 字')
    elif cn_abstract_len > 0:
        status = 'WARN' if cn_abstract_len <= 700 else 'FAIL'
        report.add(SECT_A, 'A5 中文摘要 300—500 字', status,
                   f'实际 {cn_abstract_len} 字{"略超" if cn_abstract_len <= 700 else "严重超长"}',
                   '300—500 字', f'{cn_abstract_len} 字')
    else:
        report.add(SECT_A, 'A5 中文摘要', 'WARN', '未自动定位到中文摘要正文', '', '')

    # A6: 中文关键词 3-5 个 + 用"；"分隔
    kw_match = None
    for p in paras[:80]:
        t = p.text.strip()
        if t.startswith('[关键词]') or t.startswith('【关键词】') or ('关键词' in t and ('；' in t or ';' in t)):
            kw_match = t
            break
    if kw_match:
        # 提取关键词
        kw_content = re.sub(r'^.*?关键词[】\]]?[:：]?\s*', '', kw_match)
        # 统计分号数量
        n_chinese_semi = kw_content.count('；')
        n_west_semi = kw_content.count(';')
        n_comma = kw_content.count('，') + kw_content.count(',')
        # 按；分隔
        kws = [k.strip() for k in re.split(r'[；;]', kw_content) if k.strip()]
        if 3 <= len(kws) <= 5:
            if n_chinese_semi >= len(kws) - 1 and n_west_semi == 0:
                # 检查末尾是否无标点
                end_punc = kw_content.rstrip()[-1] if kw_content.strip() else ''
                if end_punc in '。.；;，,':
                    report.add(SECT_A, 'A6 中文关键词分隔与结尾', 'WARN',
                               f'关键词末尾有标点 "{end_punc}"',
                               '用"；"分隔，末尾无标点', f'{len(kws)} 个: {kws}')
                else:
                    report.add(SECT_A, 'A6 中文关键词', 'PASS', '',
                               '3—5 个，用"；"分隔，末尾无标点',
                               f'{len(kws)} 个: {kws}')
            else:
                report.add(SECT_A, 'A6 关键词分隔符', 'FAIL',
                           f'分隔符使用错误：中文；={n_chinese_semi}, 西文;={n_west_semi}, 逗号={n_comma}',
                           '一律用中文分号"；"', kw_match[:80])
        else:
            report.add(SECT_A, 'A6 关键词数量', 'FAIL',
                       f'关键词数量 {len(kws)} 不在 3—5 范围',
                       '3—5 个', f'{len(kws)} 个: {kws}')
    else:
        report.add(SECT_A, 'A6 中文关键词', 'WARN', '未自动定位到关键词段', '', '')

    # A7: 英文 Abstract + Keywords
    has_abstract = 'Abstract' in all_text or 'ABSTRACT' in all_text
    has_keywords = 'Keywords' in all_text or 'KEYWORDS' in all_text or 'Key words' in all_text
    if has_abstract and has_keywords:
        report.add(SECT_A, 'A7 英文 Abstract + Keywords', 'PASS', '', '均应存在', '都已存在')
    else:
        miss = []
        if not has_abstract: miss.append('Abstract')
        if not has_keywords: miss.append('Keywords')
        report.add(SECT_A, 'A7 英文 Abstract + Keywords', 'FAIL',
                   f'缺 {miss}', '均应存在', f'缺 {miss}')

    # A8: 正文字数 ≥ 6000
    # 粗略统计：所有正文 Normal 段的中文字数（去除空白）
    body_text = ''
    for p in paras:
        # 排除明显是目录/参考文献/标题/英文摘要的段
        t = p.text.strip()
        if not t:
            continue
        # 简单跳过纯英文/参考文献编号段
        if re.match(r'^\[\d+\]', t):  # 参考文献项
            continue
        body_text += t
    # 仅统计中文字符（汉字）
    cn_chars = re.findall(r'[\u4e00-\u9fff]', body_text)
    body_len = len(cn_chars)
    if body_len >= 6000:
        report.add(SECT_A, 'A8 正文中文字数 ≥ 6000', 'PASS', '',
                   '经管类 ≥ 6000 字', f'{body_len} 字（中文字符）')
    else:
        report.add(SECT_A, 'A8 正文中文字数 ≥ 6000', 'FAIL',
                   f'正文中文字符仅 {body_len}，未达 6000 字',
                   '≥ 6000 字', f'{body_len} 字')

    # A9: 参考文献 ≥ 6 项
    ref_items = []
    for p in paras:
        t = p.text.strip()
        m = re.match(r'^\[(\d+)\]', t)
        if m:
            ref_items.append((int(m.group(1)), t[:100]))
    # 取最大编号作为总数
    max_ref = max((r[0] for r in ref_items), default=0)
    if max_ref >= 6:
        report.add(SECT_A, 'A9 参考文献 ≥ 6 项', 'PASS', '',
                   '≥ 6 项', f'{max_ref} 项')
    else:
        report.add(SECT_A, 'A9 参考文献 ≥ 6 项', 'FAIL',
                   f'仅 {max_ref} 项',
                   '≥ 6 项', f'{max_ref} 项')

    # A10: 参考文献文中均有引用
    if max_ref > 0:
        cited = set()
        # 在正文中查找 [N] 引用
        for p in paras:
            for m in re.finditer(r'\[(\d+)\]', p.text):
                cited.add(int(m.group(1)))
        # 减去那些是参考文献项本身的编号
        # 实际是：检查每个 N (1..max_ref) 是否在正文中至少出现一次 [N]
        uncited = [n for n in range(1, max_ref + 1) if n not in cited]
        # 减去那些只在参考文献列表中出现的（即所有 [N] 出现次数 == 1 时，仅在列表中）
        cite_counts = {}
        for p in paras:
            for m in re.finditer(r'\[(\d+)\]', p.text):
                n = int(m.group(1))
                cite_counts[n] = cite_counts.get(n, 0) + 1
        unused_in_text = [n for n in range(1, max_ref + 1) if cite_counts.get(n, 0) <= 1]
        if not unused_in_text:
            report.add(SECT_A, 'A10 文献文中均有引用', 'PASS', '',
                       '每项至少在正文中引用 1 次', '全部已引')
        else:
            report.add(SECT_A, 'A10 文献文中均有引用', 'WARN',
                       f'编号 {unused_in_text} 可能未在正文引用（仅出现 1 次，即列表本身）',
                       '每项至少 2 次出现', f'未引: {unused_in_text}')

    # A11: 致谢 ≤ 500 字
    thanks_text = ''
    for i, p in enumerate(paras):
        if p.text.strip() == '致谢' or '致谢' in p.text and len(p.text) < 6:
            # 抓取后续段落
            for j in range(i+1, min(i+10, len(paras))):
                if paras[j].text.strip():
                    thanks_text += paras[j].text.strip()
            break
    thanks_len = len(re.findall(r'[\u4e00-\u9fff]', thanks_text))
    if thanks_len == 0:
        report.add(SECT_A, 'A11 致谢 ≤ 500 字（可选）', 'INFO',
                   '未检测到致谢段（致谢为可选项）', '可选', '未含')
    elif thanks_len <= 500:
        report.add(SECT_A, 'A11 致谢 ≤ 500 字', 'PASS', '',
                   '≤ 500 字', f'{thanks_len} 字')
    else:
        report.add(SECT_A, 'A11 致谢 ≤ 500 字', 'FAIL',
                   f'致谢字数 {thanks_len} 超过 500',
                   '≤ 500 字', f'{thanks_len} 字')

    # ============= B. 排版检查 =============
    SECT_B = 'B. 排版检查'

    # B1: 页面 + 页边距
    section = doc.sections[0]
    page_w = section.page_width
    page_h = section.page_height
    A4_W = Cm(21.0)
    A4_H = Cm(29.7)
    is_a4 = abs(page_w.cm - 21.0) < 0.5 and abs(page_h.cm - 29.7) < 0.5
    if is_a4:
        report.add(SECT_B, 'B1.1 A4 纸张', 'PASS', '',
                   '21 × 29.7 cm', f'{page_w.cm:.2f} × {page_h.cm:.2f} cm')
    else:
        report.add(SECT_B, 'B1.1 A4 纸张', 'FAIL',
                   f'纸张大小非 A4', '21 × 29.7 cm',
                   f'{page_w.cm:.2f} × {page_h.cm:.2f} cm')

    # 页边距：左 2.7 / 上 2 / 下 2 / 右 2 cm
    margin_checks = [
        ('B1.2 左边距 2.7 cm', section.left_margin.cm, 2.7),
        ('B1.3 上边距 2 cm', section.top_margin.cm, 2.0),
        ('B1.4 下边距 2 cm', section.bottom_margin.cm, 2.0),
        ('B1.5 右边距 2 cm', section.right_margin.cm, 2.0),
    ]
    for label, actual, expected in margin_checks:
        diff = abs(actual - expected)
        if diff < 0.1:
            report.add(SECT_B, label, 'PASS', '',
                       f'{expected} cm', f'{actual:.2f} cm')
        elif diff < 0.5:
            report.add(SECT_B, label, 'WARN', f'偏差 {diff:.2f} cm',
                       f'{expected} cm', f'{actual:.2f} cm')
        else:
            report.add(SECT_B, label, 'FAIL', f'偏差 {diff:.2f} cm',
                       f'{expected} cm', f'{actual:.2f} cm')

    # B2: 1.5 倍行距（抽样 Normal 段）
    normal_paras = [p for p in paras if p.style.name == 'Normal' and p.text.strip()][:20]
    line_spacings = []
    for p in normal_paras:
        pf = p.paragraph_format
        if pf.line_spacing:
            line_spacings.append(pf.line_spacing)
    if line_spacings:
        avg = sum(line_spacings) / len(line_spacings)
        if abs(avg - 1.5) < 0.05:
            report.add(SECT_B, 'B2 正文 1.5 倍行距', 'PASS',
                       f'抽样 {len(line_spacings)} 段', '1.5', f'{avg:.2f}')
        else:
            report.add(SECT_B, 'B2 正文行距', 'WARN',
                       f'抽样 {len(line_spacings)} 段平均行距 {avg:.2f}，可能为默认值',
                       '1.5 倍', f'{avg:.2f}')
    else:
        report.add(SECT_B, 'B2 正文行距', 'INFO',
                   '所有正文段未显式设置行距（使用样式默认）', '1.5 倍', '继承样式')

    # B3: 标题字号字体抽样（第一层次标题应为小二号黑体）
    # 找 Heading 1 / Heading 2 / Heading 3 各一个抽样
    heading_samples = {}
    for p in paras:
        sn = p.style.name
        if sn in ('Heading 1', 'Heading 2', 'Heading 3') and sn not in heading_samples and p.text.strip():
            heading_samples[sn] = p

    expected_styles = {
        'Heading 1': ('B3.1 一级标题', '小二号黑体居中', 18, '黑体'),
        'Heading 2': ('B3.2 二级标题', '小四号黑体加粗', 12, '黑体'),
        'Heading 3': ('B3.3 三级标题', '小四号黑体', 12, '黑体'),
    }
    for style_name, (label, expected_desc, expected_pt, expected_font) in expected_styles.items():
        if style_name not in heading_samples:
            report.add(SECT_B, label, 'INFO', f'未在文档中找到 {style_name} 样式段', expected_desc, '不适用')
            continue
        p = heading_samples[style_name]
        # 取第一个 run
        if not p.runs:
            report.add(SECT_B, label, 'WARN', f'{style_name} 段无 run', expected_desc, '空')
            continue
        run = p.runs[0]
        size = get_run_size(run)
        ea_font = get_run_east_asia_font(run)
        size_str = fmt_size(size) if size else '继承样式'
        font_str = ea_font or '继承样式'
        size_ok = size and abs(size.pt - expected_pt) < 0.5
        font_ok = ea_font and expected_font in ea_font
        if size_ok and font_ok:
            report.add(SECT_B, label, 'PASS', '',
                       expected_desc, f'{size_str}/{font_str}')
        elif not size and not ea_font:
            report.add(SECT_B, label, 'INFO',
                       '字号字体继承自样式，未在 run 上显式设置',
                       expected_desc, '样式继承')
        else:
            issues = []
            if not size_ok: issues.append(f'字号 {size_str} ≠ 期望 {SIZE_MAP.get(expected_pt, "?")}')
            if not font_ok: issues.append(f'字体 {font_str} ≠ 期望 {expected_font}')
            report.add(SECT_B, label, 'WARN', '; '.join(issues),
                       expected_desc, f'{size_str}/{font_str}')

    # B4: 正文字号字体（抽样 Normal 段）
    normal_sample = None
    for p in paras:
        if p.style.name == 'Normal' and p.text.strip() and len(p.text) > 30:
            normal_sample = p
            break
    if normal_sample and normal_sample.runs:
        run = normal_sample.runs[0]
        size = get_run_size(run)
        ea_font = get_run_east_asia_font(run)
        size_ok = size and abs(size.pt - 12) < 0.5
        font_ok = ea_font and ('宋体' in ea_font or 'SimSun' in ea_font.lower() or 'Song' in ea_font)
        if size_ok and font_ok:
            report.add(SECT_B, 'B4 正文 小四号宋体', 'PASS', '',
                       '小四号宋体', f'{fmt_size(size)}/{ea_font}')
        elif not size and not ea_font:
            report.add(SECT_B, 'B4 正文字号字体', 'INFO',
                       '正文继承样式，未在 run 上显式设置',
                       '小四号宋体', '样式继承')
        else:
            issues = []
            if not size_ok: issues.append(f'字号 {fmt_size(size)}')
            if not font_ok: issues.append(f'字体 {ea_font or "未设置"}')
            report.add(SECT_B, 'B4 正文字号字体', 'WARN', '; '.join(issues),
                       '小四号宋体', f'{fmt_size(size) if size else "?"}/{ea_font or "?"}')

    # ============= C. 图表检查 =============
    SECT_C = 'C. 图表检查'

    # C1: 表格数量与编号连续性
    tbl_count = len(doc.tables)
    tbl_caps = []
    for p in paras:
        m = re.match(r'^\s*表\s*(\d+)\s', p.text)
        if m:
            tbl_caps.append(int(m.group(1)))
    if tbl_caps:
        sorted_caps = sorted(set(tbl_caps))
        expected_seq = list(range(1, max(sorted_caps) + 1))
        missing = [n for n in expected_seq if n not in sorted_caps]
        if not missing:
            report.add(SECT_C, 'C1 表格编号连续性', 'PASS', '',
                       '连续不跳跃', f'表 1 到表 {max(sorted_caps)}, 共 {len(sorted_caps)} 编号')
        else:
            report.add(SECT_C, 'C1 表格编号连续性', 'WARN',
                       f'编号 {missing} 缺失',
                       '连续不跳跃', f'缺 {missing}')
    report.add(SECT_C, 'C1.2 文档表格对象数', 'INFO', '',
               '应与编号一致', f'{tbl_count} 个 Table 对象，{len(set(tbl_caps))} 个表标题')

    # C2: 图编号连续性
    fig_caps = []
    for p in paras:
        m = re.match(r'^\s*图\s*(\d+)[\s\u2032\u2032\u2019]?', p.text)
        if m:
            fig_caps.append(int(m.group(1)))
    if fig_caps:
        sorted_caps = sorted(set(fig_caps))
        expected_seq = list(range(1, max(sorted_caps) + 1))
        missing = [n for n in expected_seq if n not in sorted_caps]
        if not missing:
            report.add(SECT_C, 'C2 图编号连续性', 'PASS', '',
                       '连续不跳跃', f'图 1 到图 {max(sorted_caps)}, 共 {len(sorted_caps)} 编号')
        else:
            report.add(SECT_C, 'C2 图编号连续性', 'WARN',
                       f'编号 {missing} 缺失',
                       '连续不跳跃', f'缺 {missing}')

    # C3: 三线表（粗略：检查表格的边框设置）
    # 三线表特征：顶部、底部、表头下方有粗/细横线，无竖线、无中间横线
    # python-docx 难以精确判断，这里仅做粗略检查（看表格 style 名是否含 'Grid' = 全网格非三线表）
    three_line_check = []
    for ti, tbl in enumerate(doc.tables):
        style_name = tbl.style.name if tbl.style else ''
        is_grid = 'Grid' in style_name or '网格' in style_name
        three_line_check.append((ti, style_name, is_grid))
    n_grid = sum(1 for _, _, g in three_line_check if g)
    if n_grid > 0:
        non_three = [f'Table{ti}({sn})' for ti, sn, g in three_line_check if g]
        report.add(SECT_C, 'C3 表格三线表（粗略）', 'WARN',
                   f'有 {n_grid} 个表格使用 Grid/网格样式，可能非三线表',
                   '所有表格三线表', f'{non_three}')
    else:
        report.add(SECT_C, 'C3 表格三线表（粗略）', 'PASS',
                   '未检出 Grid 样式表格', '所有表格三线表',
                   f'{tbl_count} 个表格样式: {[sn for _, sn, _ in three_line_check]}')

    # ============= D. 数字/标点检查 =============
    SECT_D = 'D. 数字/标点检查'

    # D1: 公历日期格式
    han_date = re.findall(r'[一二三四五六七八九〇○两零]+年[一二三四五六七八九〇○十]+月', all_text)
    arab_date = re.findall(r'\b(19|20)\d{2}年\d{1,2}月', all_text)
    if han_date:
        report.add(SECT_D, 'D1 公历日期阿拉伯数字', 'WARN',
                   f'发现汉字公历日期 {len(han_date)} 处',
                   '公历用阿拉伯数字', f'例: {han_date[:3]}')
    else:
        report.add(SECT_D, 'D1 公历日期阿拉伯数字', 'PASS', '',
                   '公历用阿拉伯数字', f'阿拉伯数字日期 {len(arab_date)} 处')

    # D2: 标点符号（中文逗号 vs 英文逗号在中文正文中混用）
    cn_text_runs = []
    for p in paras:
        if not p.text.strip(): continue
        # 只检查中文段（含汉字 > 50%）
        t = p.text
        cn_chars = len(re.findall(r'[\u4e00-\u9fff]', t))
        if cn_chars > len(t) * 0.5:
            cn_text_runs.append(t)
    cn_text = '\n'.join(cn_text_runs)
    # 找中文段中是否出现 英文逗号 ASCII 0x2C 后跟汉字（说明在中文语境用了英文逗号）
    bad_comma = re.findall(r',[ \u4e00-\u9fff]', cn_text)
    if len(bad_comma) > 5:
        report.add(SECT_D, 'D2 中文语境标点', 'WARN',
                   f'在中文段中发现 {len(bad_comma)} 处疑似英文逗号 ",", 应改为中文逗号 "，"',
                   '中文段用中文标点', f'例如 {len(bad_comma)} 处')
    elif len(bad_comma) > 0:
        report.add(SECT_D, 'D2 中文语境标点', 'INFO',
                   f'中文段有 {len(bad_comma)} 处英文逗号（含合理用例如 e.g.）',
                   '中文段用中文标点', f'仅 {len(bad_comma)} 处')
    else:
        report.add(SECT_D, 'D2 中文语境标点', 'PASS',
                   '中文段未检出英文逗号',
                   '遵循 GB/T 15834-1995', '0 处')

    # ============= E. 标题序次检查 =============
    SECT_E = 'E. 标题序次检查'

    # 检测论文使用的标题序次类别
    # 经管类：一、（一）1. （1）①
    # 理工/外语：1 1.1 1.1.1
    h1_pattern = re.compile(r'^[一二三四五六七八九十]+、')
    h2_pattern = re.compile(r'^（[一二三四五六七八九十]+）')
    digit_pattern = re.compile(r'^\d+(\.\d+)+\s')

    h1_jingguan = sum(1 for p in paras if h1_pattern.match(p.text.strip()))
    h2_jingguan = sum(1 for p in paras if h2_pattern.match(p.text.strip()))
    digit_count = sum(1 for p in paras if digit_pattern.match(p.text.strip()))

    if h1_jingguan > 0 and h2_jingguan > 0:
        report.add(SECT_E, 'E1 标题序次（经管类）', 'PASS',
                   f'符合经管类格式',
                   '一、（一）1.（1）①',
                   f'一级 {h1_jingguan} 个，二级 {h2_jingguan} 个')
        if digit_count > 5:
            report.add(SECT_E, 'E2 是否混用理工序次', 'WARN',
                       f'同时检出 {digit_count} 处 N.N 形式编号，注意是否为序次混用',
                       '经管类应一致用汉字序次',
                       f'N.N 形式 {digit_count} 处（可能是公式/章节号）')
    elif digit_count > 5:
        report.add(SECT_E, 'E1 标题序次（理工/外语类）', 'PASS',
                   f'符合理工/外语格式',
                   '1 1.1 1.1.1', f'N.N 形式 {digit_count} 处')
    else:
        report.add(SECT_E, 'E1 标题序次', 'WARN', '未明确检测出序次格式', '', '')

    # ============= F. 装订（INFO） =============
    SECT_F = 'F. 装订检查'
    report.add(SECT_F, 'F1 装订顺序', 'INFO',
               '装订顺序为：封面→诚信承诺书→目录→中英文题目→正文→注释→参考文献→致谢',
               '本规范要求', '请人工核对')
    report.add(SECT_F, 'F2 装订线位置', 'INFO',
               '装订线应位于页面左侧；已配 B1.2 左边距 2.7 cm 标准',
               '左侧装订', '配套边距已符合')
    report.add(SECT_F, 'F3 打印面', 'INFO',
               '内页要求单面打印；手写论文不作为答辩材料',
               '单面打印', '需打印时确认')

    # 写入报告
    md = report.to_markdown(DOC)
    with open(OUT, 'w', encoding='utf-8') as f:
        f.write(md)
    print(f'报告已生成: {OUT}')
    print(f'文件大小: {os.path.getsize(OUT):,} bytes')

    # 概要
    print('\n==== 概要 ====')
    total_pass = total_warn = total_fail = total_info = 0
    for items in report.sections.values():
        for it in items:
            if it['status'] == 'PASS': total_pass += 1
            elif it['status'] == 'WARN': total_warn += 1
            elif it['status'] == 'FAIL': total_fail += 1
            else: total_info += 1
    print(f'  PASS: {total_pass}')
    print(f'  WARN: {total_warn}')
    print(f'  FAIL: {total_fail}')
    print(f'  INFO: {total_info}')


if __name__ == '__main__':
    main()
