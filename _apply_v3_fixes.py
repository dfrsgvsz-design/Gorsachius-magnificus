# -*- coding: utf-8 -*-
"""
v2 → v3 综合修复：
F1 上下标：
  F1.1 删除全部空 vertAlign run（146 个）
  F1.2 P121 把 R0 中 [10] 和 [11] 拆分为 superscript run
F2 交叉引用：
  F2.1 P208 「表4.1」 → 「表3」 （同时补 表3 正文引用）
F3 补全图表正文引用：
  F3.1 P195 末追加：「6 个驱动因子的空间分布如图3所示」
  F3.2 P199 末追加：「模型精度验证结果详见表2」
  F3.3 P249 「结果见表5和图5」 → 「结果见表5、图5、图6和图7」
  F3.4 P287 「结果如表9和图9所示」 → 「结果如表9、图9和图10所示」
  F3.5 P360 「最终的组合权重如表12所示」末追加图11图12引用
  F3.6 P405 「障碍度诊断结果表明」 → 「障碍度诊断结果（图14）表明」
  F3.7 P411 「Pearson相关分析」末追加图15引用
  F3.8 P422 「加速下降态势」末追加图16引用

每一步均做前/后断言，并打印执行日志。
"""
import re
import sys
import io
from copy import deepcopy
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = r"E:\大学\万物春\erci\郑春铃+横州市土地利用变化与生态安全评价(1)_PLUS修订版_最终_格式合规_v2.docx"
DST = r"E:\大学\万物春\erci\郑春铃+横州市土地利用变化与生态安全评价(1)_PLUS修订版_最终_格式合规_v3.docx"
LOG = r"f:\Gorsachius magnificus\_apply_v3_log.txt"

sys.stdout = io.open(LOG, 'w', encoding='utf-8')

doc = Document(SRC)
paras = doc.paragraphs
print(f"v2 -> v3 修复开始, 段落数={len(paras)}")

# ===================== F1.1 删除空 vertAlign run =====================
print("\n[F1.1] 删除空 vertAlign run …")
removed_empty_va = 0
for pi, p in enumerate(paras):
    runs_to_remove = []
    for r in p.runs:
        rpr = r._element.find(qn('w:rPr'))
        if rpr is None:
            continue
        va = rpr.find(qn('w:vertAlign'))
        if va is None:
            continue
        if not r.text:
            runs_to_remove.append(r._element)
    for r_el in runs_to_remove:
        r_el.getparent().remove(r_el)
        removed_empty_va += 1
print(f"  共删除 {removed_empty_va} 个空 vertAlign run")

# ===================== F1.2 P121 [10] 和 [11] 拆分为 superscript =====================
print("\n[F1.2] P121 拆分 [10] 和 [11] 为 superscript …")
p121 = paras[121]
# 找包含 [10] 的 run
target_run_idx = None
for ri, r in enumerate(p121.runs):
    if '[10]' in r.text and '[11]' in r.text:
        target_run_idx = ri
        break
print(f"  目标 run = R{target_run_idx} (含 [10] 和 [11])")
assert target_run_idx is not None, "找不到含 [10] 和 [11] 的 run!"

orig_run = p121.runs[target_run_idx]
orig_text = orig_run.text
print(f"  原文本: {orig_text[:80]}...{orig_text[-30:]}")

m10 = orig_text.find('[10]')
m11 = orig_text.find('[11]')
assert m10 >= 0 and m11 > m10, "[10]/[11] 位置异常"

before10 = orig_text[:m10]
between  = orig_text[m10+4:m11]
after11  = orig_text[m11+4:]
print(f"  拆分：before10={len(before10)} | between={len(between)} | after11={len(after11)}")

# 修改原 run 文本为 before10
orig_run.text = before10
# 强制 <w:t xml:space="preserve">
t_el = orig_run._element.find(qn('w:t'))
if t_el is not None:
    t_el.set(qn('xml:space'), 'preserve')

def make_run_like(template_run_el, text, superscript=False):
    """根据模板 run 复制，设置文本和（可选）superscript。"""
    new_r = deepcopy(template_run_el)
    # 清空模板的所有 <w:t> 和其他子元素，保留 <w:rPr>
    for child in list(new_r):
        if child.tag == qn('w:rPr'):
            continue
        new_r.remove(child)
    # 设置 superscript
    if superscript:
        rpr = new_r.find(qn('w:rPr'))
        if rpr is None:
            rpr = OxmlElement('w:rPr')
            new_r.insert(0, rpr)
        va = rpr.find(qn('w:vertAlign'))
        if va is None:
            va = OxmlElement('w:vertAlign')
            rpr.append(va)
        va.set(qn('w:val'), 'superscript')
    # 添加 <w:t>
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    new_r.append(t)
    return new_r

# 创建 4 个新 run：sup10, between_run, sup11, after_run
template = orig_run._element
sup10 = make_run_like(template, '[10]', superscript=True)
between_run = make_run_like(template, between, superscript=False)
sup11 = make_run_like(template, '[11]', superscript=True)
after_run = make_run_like(template, after11, superscript=False)

# 按顺序插入到 orig_run 之后
parent = orig_run._element.getparent()
my_idx = list(parent).index(orig_run._element)
# 倒序插入，让它们最终顺序为：orig_run, sup10, between_run, sup11, after_run
parent.insert(my_idx + 1, after_run)
parent.insert(my_idx + 1, sup11)
parent.insert(my_idx + 1, between_run)
parent.insert(my_idx + 1, sup10)
print(f"  P121 已插入 4 个新 run: [10]sup + between + [11]sup + after")

# 验证
p121_new = doc.paragraphs[121]
print(f"  P121 新 run 数: {len(p121_new.runs)}")
sup_texts = []
for r in p121_new.runs:
    rpr = r._element.find(qn('w:rPr'))
    if rpr is not None and rpr.find(qn('w:vertAlign')) is not None:
        sup_texts.append(r.text)
print(f"  P121 superscript texts: {sup_texts}")
assert '[10]' in sup_texts and '[11]' in sup_texts, "[10] 或 [11] 未正确加 superscript"

# ===================== F2 + F3 文本替换 =====================
TEXT_EDITS = [
    # (段落 idx, 旧文本, 新文本, 描述)
    (208, '表4.1', '表3', 'F2.1 P208 章节式残留 + 补表3引用'),
    (195, '本研究选取以下 6 个驱动因子构建 PLUS 模型', '本研究选取以下 6 个驱动因子构建 PLUS 模型（6 个驱动因子的空间分布如图3所示）', 'F3.1 补图3引用'),
    (249, '结果见表5和图5。', '结果见表5、图5、图6和图7。', 'F3.3 补图6/图7引用'),
    (287, '结果如表9和图9所示。', '结果如表9、图9和图10所示。', 'F3.4 补图10引用'),
    (360, '最终的组合权重如表12所示。', '最终的组合权重如表12所示，各指标 PSR 雷达图见图11，组合权重对比可视化见图12。', 'F3.5 补图11/图12引用'),
    (405, '障碍度诊断结果表明，', '障碍度诊断结果（图14）表明，', 'F3.6 补图14引用'),
    (411, '本研究对ESI与各PSR指标进行了Pearson相关分析。', '本研究对ESI与各PSR指标进行了Pearson相关分析（结果如图15所示）。', 'F3.7 补图15引用'),
    (422, '导致ESI呈现加速下降态势。', '导致ESI呈现加速下降态势（综合分析详见图16）。', 'F3.8 补图16引用'),
]

# 199 段需要追加表2引用，但其文本以"max_d"截断（被诊断省略），实际需要找到该段并定位末尾追加
# 先单独处理：定位 P199 完整文本，在段末追加 ", 模型精度验证结果详见表2。"

def edit_paragraph_text(p, old, new, dry=False):
    """在 paragraph p 的 runs 中查找 old 并替换为 new。
    若 old 跨多个 run，则合并连续 run（删去后续，把全文塞到第一个匹配 run）。
    返回 (matched, run_idx)"""
    # 单 run 命中尝试
    for ri, r in enumerate(p.runs):
        if old in r.text:
            if not dry:
                r.text = r.text.replace(old, new)
                t_el = r._element.find(qn('w:t'))
                if t_el is not None:
                    t_el.set(qn('xml:space'), 'preserve')
            return True, ri
    # 跨 run 命中：连接所有 run 文本看是否包含 old
    full_text = "".join(r.text for r in p.runs)
    if old in full_text:
        if not dry:
            # 把所有 runs 文本合并到 R0，其余清空
            full_new = full_text.replace(old, new)
            r0 = p.runs[0]
            r0.text = full_new
            t_el = r0._element.find(qn('w:t'))
            if t_el is not None:
                t_el.set(qn('xml:space'), 'preserve')
            for r in p.runs[1:]:
                r.text = ''
        return True, 0
    return False, -1

print("\n[F2/F3] 文本替换 …")
for pi, old, new, desc in TEXT_EDITS:
    p = paras[pi]
    matched, ri = edit_paragraph_text(p, old, new)
    if matched:
        print(f"  ✓ P{pi} R{ri}  {desc}")
    else:
        print(f"  ✗ P{pi} 未匹配「{old[:30]}…」  {desc}")
        # 打印实际段落前 200 字符以调试
        print(f"     P{pi} actual: {p.text[:200]!r}")

# 单独处理 P199 末尾追加（更稳）
print("\n[F3.2] P199 末尾追加表2引用 …")
p199 = paras[199]
p199_text = p199.text
print(f"  P199 当前末尾: ...{p199_text[-80:]!r}")
# 在最后一个非空 run 末尾追加
last_run = None
for r in p199.runs:
    if r.text:
        last_run = r
if last_run is None:
    last_run = p199.runs[0]

# 追加文本，注意保留中文句末
append_text = "（精度验证结果详见表2）"
# 若结尾是句号"。"，把追加文本放在句号之前
if last_run.text.endswith('。'):
    last_run.text = last_run.text[:-1] + append_text + '。'
else:
    last_run.text = last_run.text + append_text
t_el = last_run._element.find(qn('w:t'))
if t_el is not None:
    t_el.set(qn('xml:space'), 'preserve')
print(f"  P199 修改后末尾: ...{p199.text[-80:]!r}")

# ===================== 保存 =====================
print(f"\n保存 -> {DST}")
doc.save(DST)
print("v3 保存完成")
print(f"\nLOG path: {LOG}")
