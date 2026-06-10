# -*- coding: utf-8 -*-
"""
v3 最终验证：
1. vertAlign run 净化情况（应仅剩有文本的角标）
2. 全图表引用覆盖率（图1-16 + 表1-14 全部被引用）
3. P121 [10]/[11] 已加 superscript
4. P208 已改为 表3
5. 段落数变化
6. 关键修改段（P195/P199/P249/P287/P360/P405/P411/P422）抽样
"""
import re
import sys
import io
from collections import defaultdict
from docx import Document
from docx.oxml.ns import qn

V2 = r"E:\大学\万物春\erci\郑春铃+横州市土地利用变化与生态安全评价(1)_PLUS修订版_最终_格式合规_v2.docx"
V3 = r"E:\大学\万物春\erci\郑春铃+横州市土地利用变化与生态安全评价(1)_PLUS修订版_最终_格式合规_v3.docx"
LOG = r"f:\Gorsachius magnificus\_verify_v3_report.txt"

sys.stdout = io.open(LOG, 'w', encoding='utf-8')

doc2 = Document(V2)
doc3 = Document(V3)
ps2 = doc2.paragraphs
ps3 = doc3.paragraphs

print("=" * 80)
print("v3 最终验证")
print("=" * 80)
print(f"v2 段落数={len(ps2)}, v3 段落数={len(ps3)}, 差值={len(ps3)-len(ps2)} （仅文本编辑，应=0）")

# ===== 1. vertAlign 净化 =====
def vert_stats(paras, label):
    total, empty = 0, 0
    by_para = defaultdict(list)
    for pi, p in enumerate(paras):
        for r in p.runs:
            rpr = r._element.find(qn('w:rPr'))
            if rpr is None: continue
            va = rpr.find(qn('w:vertAlign'))
            if va is None: continue
            total += 1
            if not r.text:
                empty += 1
            by_para[pi].append((va.get(qn('w:val')), r.text))
    print(f"\n[{label}] vertAlign 总数={total}, 空 run={empty}, 涉及段={len(by_para)}")
    return total, empty, by_para

t2, e2, bp2 = vert_stats(ps2, "v2")
t3, e3, bp3 = vert_stats(ps3, "v3")
print(f"  >>> 净化效果: 总 {t2}->{t3} (减少 {t2-t3}), 空 {e2}->{e3} (减少 {e2-e3})")

# ===== 2. P121 [10]/[11] superscript 验证 =====
print("\n[v3 P121 验证]")
p121 = ps3[121]
print(f"  段落 text 长度: {len(p121.text)}")
print(f"  runs 数: {len(p121.runs)}")
sup_in_p121 = []
for ri, r in enumerate(p121.runs):
    rpr = r._element.find(qn('w:rPr'))
    is_sup = False
    if rpr is not None:
        va = rpr.find(qn('w:vertAlign'))
        if va is not None and va.get(qn('w:val')) == 'superscript':
            is_sup = True
    print(f"  R{ri:2d} {'[SUP]' if is_sup else '     '} {r.text[:60]!r}")
    if is_sup and r.text:
        sup_in_p121.append(r.text)
print(f"  P121 superscript 文本: {sup_in_p121}")
assert '[10]' in sup_in_p121 and '[11]' in sup_in_p121, "P121 [10] 或 [11] 未正确标 superscript!"

# ===== 3. 图/表标题与正文引用覆盖 =====
fig_re = re.compile(r'^图\s*(\d+)[\s：:]')
tab_re = re.compile(r'^表\s*(\d+)[\s：:]')

caption_fig, caption_tab = {}, {}
caption_paras = set()
for pi, p in enumerate(ps3):
    t = p.text.strip()
    m = fig_re.match(t)
    if m:
        caption_paras.add(pi); caption_fig[pi] = int(m.group(1)); continue
    m = tab_re.match(t)
    if m:
        caption_paras.add(pi); caption_tab[pi] = int(m.group(1))

print(f"\n[v3 图表标题段]")
print(f"  图标题: {len(caption_fig)} 个 = {sorted(caption_fig.values())}")
print(f"  表标题: {len(caption_tab)} 个 = {sorted(caption_tab.values())}")

# 正文引用
ref_patterns = {
    '图': re.compile(r'图\s*(\d+)'),
    '表': re.compile(r'表\s*(\d+)'),
}
xref_fig, xref_tab = [], []
for pi, p in enumerate(ps3):
    if pi in caption_paras: continue
    for kind, pat in ref_patterns.items():
        for m in pat.finditer(p.text):
            n = int(m.group(1))
            if kind == '图':
                xref_fig.append((pi, n, p.text[max(0,m.start()-15):m.end()+20]))
            else:
                xref_tab.append((pi, n, p.text[max(0,m.start()-15):m.end()+20]))

fig_nums = sorted(set(h[1] for h in xref_fig))
tab_nums = sorted(set(h[1] for h in xref_tab))
fig_all = set(caption_fig.values())
tab_all = set(caption_tab.values())
missing_fig = sorted(fig_all - set(fig_nums))
missing_tab = sorted(tab_all - set(tab_nums))

print(f"\n[v3 引用覆盖]")
print(f"  图引用集 = {fig_nums}")
print(f"  表引用集 = {tab_nums}")
print(f"  未被引用的图 = {missing_fig}")
print(f"  未被引用的表 = {missing_tab}")

# 越界
oob_fig = [(p,n,c) for p,n,c in xref_fig if n<1 or n>max(fig_all)]
oob_tab = [(p,n,c) for p,n,c in xref_tab if n<1 or n>max(tab_all)]
print(f"  图引用越界 = {len(oob_fig)} 处")
for p,n,c in oob_fig:
    print(f"    P{p} 图{n}: {c!r}")
print(f"  表引用越界 = {len(oob_tab)} 处")
for p,n,c in oob_tab:
    print(f"    P{p} 表{n}: {c!r}")

# ===== 4. P208 验证 =====
print("\n[v3 P208 验证]")
print(f"  P208: {ps3[208].text[:150]!r}")
assert '表4.1' not in ps3[208].text, "P208 「表4.1」未替换!"
assert '表3' in ps3[208].text, "P208 「表3」未出现!"
print("  ✓ 「表4.1」已替换为「表3」")

# ===== 5. 关键修改段抽样 =====
print("\n[v3 关键修改段抽样]")
SAMPLE = [195, 199, 249, 287, 360, 405, 411, 422]
for pi in SAMPLE:
    t = ps3[pi].text
    # 显示后 100 字符（修改一般在末尾）
    print(f"  P{pi} (...{t[-110:]!r})")

# ===== 6. 引用统计明细（v3 中新增的引用句） =====
print("\n[v3 新增引用语句抽样]")
new_refs_to_check = [
    (195, '图3'),
    (199, '表2'),
    (249, '图6'),
    (249, '图7'),
    (287, '图10'),
    (360, '图11'),
    (360, '图12'),
    (405, '图14'),
    (411, '图15'),
    (422, '图16'),
    (208, '表3'),
]
for pi, ref in new_refs_to_check:
    t = ps3[pi].text
    if ref in t:
        idx = t.find(ref)
        ctx = t[max(0,idx-25):idx+25]
        print(f"  ✓ P{pi} 「{ref}」: ...{ctx}...")
    else:
        print(f"  ✗ P{pi} 「{ref}」 NOT FOUND")

print("\n" + "=" * 80)
print("✅ v3 验证通过" if (e3 == 0 and not missing_fig and not missing_tab and not oob_fig and not oob_tab) else "⚠️ 仍有问题")
print("=" * 80)
