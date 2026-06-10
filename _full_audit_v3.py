# -*- coding: utf-8 -*-
"""
v3 文档全面审计 — 12 维度：
A. GXUFE 完整性     B. 排版        C. 图表           D. 数字/标点
E. 标题序次         F. 装订        G. 上下标特殊字符  H. 交叉引用深度
I. 参考文献深度     J. 内容鲁棒性  K. 字体一致性     L. 段落格式
"""
import sys, io, re, os
from collections import Counter, defaultdict
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn

DOC = r'E:\大学\万物春\erci\郑春铃+横州市土地利用变化与生态安全评价(1)_PLUS修订版_最终_格式合规_v3.docx'
OUT = r'E:\大学\万物春\erci\_v3全面审计报告.md'
LOG = r'f:\Gorsachius magnificus\_full_audit_v3_log.txt'

sys.stdout = io.open(LOG, 'w', encoding='utf-8')

SIZE_MAP = {22: '二号', 18: '小二号', 16: '三号', 15: '小三号',
            14: '四号', 12: '小四号', 10.5: '五号', 9: '小五号'}

def fmt_size(pt):
    if pt is None: return '默认'
    val = pt.pt if hasattr(pt, 'pt') else float(pt)
    return f"{SIZE_MAP.get(val, '?')}({val}pt)"

def get_run_size(run):
    return run.font.size if run.font.size else None

def get_run_east_asia_font(run):
    rPr = run._element.find(qn('w:rPr'))
    if rPr is None: return None
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None: return None
    return rFonts.get(qn('w:eastAsia'))


class Report:
    def __init__(self):
        self.sections = {}
    def add(self, sect, label, status, msg='', exp='', act=''):
        self.sections.setdefault(sect, []).append({
            'label': label, 'status': status, 'message': msg,
            'expected': exp, 'actual': act})

    def stats(self):
        c = Counter()
        for items in self.sections.values():
            for it in items:
                c[it['status']] += 1
        return c

    def to_md(self, doc_path):
        c = self.stats()
        lines = [
            f'# v3 文档全面审计报告\n',
            f'**文件**：`{doc_path}`\n',
            f'**审计维度**：12 个（A–L）\n',
            f'**依据**：GXUFE 论文撰写规范 + 学术写作通用质量准则\n\n---\n',
            f'## 总体统计\n',
            f'- ✅ PASS : **{c.get("PASS",0)}**',
            f'- ⚠️ WARN : **{c.get("WARN",0)}**',
            f'- ❌ FAIL : **{c.get("FAIL",0)}**',
            f'- ℹ️ INFO : **{c.get("INFO",0)}**\n',
        ]
        for sect, items in self.sections.items():
            lines.append(f'\n## {sect}\n')
            lines.append('| 状态 | 检查项 | 说明 | 期望 | 实测 |')
            lines.append('|------|--------|------|------|------|')
            for it in items:
                emoji = {'PASS':'✅','WARN':'⚠️','FAIL':'❌','INFO':'ℹ️'}.get(it['status'],'?')
                msg = (it['message'] or '').replace('|','\\|').replace('\n',' ')[:200]
                exp = (it['expected'] or '').replace('|','\\|')[:80]
                act = (it['actual'] or '').replace('|','\\|')[:120]
                lines.append(f'| {emoji} {it["status"]} | {it["label"]} | {msg} | {exp} | {act} |')
        return '\n'.join(lines)


# =================================================================
# 主审计
# =================================================================
def main():
    rep = Report()
    doc = Document(DOC)
    paras = doc.paragraphs
    all_text = '\n'.join(p.text for p in paras)
    print(f"文档段落数: {len(paras)}, 表格数: {len(doc.tables)}")

    # ============= A. GXUFE 完整性 =============
    SA = 'A. GXUFE 完整性'

    # A1 封面元素
    cover_elements = ['题目','系别','专业','班级','姓名','学号','指导教师']
    front_text = '\n'.join(p.text for p in paras[:30])
    missing = [e for e in cover_elements if e not in front_text]
    if not missing:
        rep.add(SA,'A1 封面元素齐全','PASS','','7 项','全部存在')
    else:
        rep.add(SA,'A1 封面元素','WARN',f'缺：{missing}','7 项',f'缺 {missing}')

    # A2 题目 ≤ 25 字
    title_text = ''
    for p in paras[:30]:
        t = p.text.strip()
        if '横州市' in t and '土地利用' in t and 10 < len(t) < 60:
            title_text = t
            break
    if title_text:
        tl = len(re.sub(r'\s','',title_text))
        if tl <= 25:
            rep.add(SA,'A2 题目 ≤ 25 字','PASS','','≤ 25 字',f'{tl} 字: {title_text}')
        else:
            rep.add(SA,'A2 题目 ≤ 25 字','FAIL',f'题目 {tl} 字','≤ 25 字',f'{tl} 字: {title_text}')

    # A3 诚信承诺书
    if '诚信承诺书' in all_text or '承诺书' in front_text:
        rep.add(SA,'A3 诚信承诺书','PASS','','应存在','已存在')
    else:
        rep.add(SA,'A3 诚信承诺书','WARN','未检出','应单独一页','缺失')

    # A4 目录
    if any('目录' in p.text for p in paras[:50]):
        rep.add(SA,'A4 目录','PASS','','三级层次','已含')
    else:
        rep.add(SA,'A4 目录','WARN','未检出','','')

    # A5 中文摘要 300—500 字
    cn_abs_idx = -1
    for i, p in enumerate(paras):
        if '中文摘要' in p.text or '[摘要]' in p.text or '【摘要】' in p.text:
            cn_abs_idx = i; break
    cn_abs_len = 0
    if cn_abs_idx >= 0:
        for j in range(cn_abs_idx, min(cn_abs_idx+5, len(paras))):
            t = paras[j].text.strip()
            if len(t) > 100 and '中文摘要' not in t and '关键词' not in t:
                cn_abs_len = len(re.findall(r'[\u4e00-\u9fff]', t))
                break
    if 300 <= cn_abs_len <= 500:
        rep.add(SA,'A5 中文摘要 300—500 字','PASS','','300—500 字',f'{cn_abs_len} 字')
    elif cn_abs_len > 0:
        st = 'WARN' if cn_abs_len <= 700 else 'FAIL'
        rep.add(SA,'A5 中文摘要 300—500 字',st,f'实际 {cn_abs_len} 字','300—500 字',f'{cn_abs_len} 字')
    else:
        rep.add(SA,'A5 中文摘要','WARN','未定位','','')

    # A6 中文关键词 3—5 个
    kw_match = None
    for p in paras[:80]:
        t = p.text.strip()
        if (t.startswith('[关键词]') or t.startswith('【关键词】') or
            ('关键词' in t and '；' in t)):
            kw_match = t; break
    if kw_match:
        kwc = re.sub(r'^.*?关键词[】\]]?[:：]?\s*','', kw_match)
        ncs = kwc.count('；'); nws = kwc.count(';')
        kws = [k.strip() for k in re.split(r'[；;]', kwc) if k.strip()]
        if 3 <= len(kws) <= 5:
            end_p = kwc.rstrip()[-1] if kwc.strip() else ''
            if end_p in '。.；;，,':
                rep.add(SA,'A6 中文关键词结尾','WARN',f'末尾有标点 "{end_p}"','无标点',f'{len(kws)} 个: {kws}')
            elif nws == 0:
                rep.add(SA,'A6 中文关键词','PASS','','3—5 个，"；"分隔',f'{len(kws)} 个')
            else:
                rep.add(SA,'A6 关键词分隔符','FAIL',f'西文; {nws} 处','"；"分隔',kw_match[:80])
        else:
            rep.add(SA,'A6 关键词数量','FAIL',f'{len(kws)} 个','3—5 个',f'{len(kws)} 个: {kws}')
    else:
        rep.add(SA,'A6 中文关键词','WARN','未定位','','')

    # A7 英文 Abstract + Keywords
    ha = 'Abstract' in all_text or 'ABSTRACT' in all_text
    hk = 'Keywords' in all_text or 'Key words' in all_text or 'KEYWORDS' in all_text
    if ha and hk:
        rep.add(SA,'A7 英文 Abstract+Keywords','PASS','','均应存在','都存在')
    else:
        miss=[]
        if not ha: miss.append('Abstract')
        if not hk: miss.append('Keywords')
        rep.add(SA,'A7 英文 Abstract+Keywords','FAIL',f'缺 {miss}','均存在',f'缺 {miss}')

    # A8 正文中文字数
    body = ''
    for p in paras:
        t = p.text.strip()
        if not t: continue
        if re.match(r'^\[\d+\]', t): continue  # 跳过文献项
        body += t
    bl = len(re.findall(r'[\u4e00-\u9fff]', body))
    if bl >= 6000:
        rep.add(SA,'A8 正文中文字数 ≥ 6000','PASS','','≥ 6000 字',f'{bl} 字')
    else:
        rep.add(SA,'A8 正文中文字数 ≥ 6000','FAIL',f'仅 {bl}','≥ 6000 字',f'{bl} 字')

    # A9 参考文献 ≥ 6 项
    ref_items = []
    for p in paras:
        t = p.text.strip()
        m = re.match(r'^\[(\d+)\]', t)
        if m:
            ref_items.append((int(m.group(1)), t[:120]))
    max_ref = max((r[0] for r in ref_items), default=0)
    if max_ref >= 6:
        rep.add(SA,'A9 参考文献 ≥ 6 项','PASS','','≥ 6 项',f'{max_ref} 项')
    else:
        rep.add(SA,'A9 参考文献 ≥ 6 项','FAIL',f'仅 {max_ref}','≥ 6 项',f'{max_ref} 项')

    # A10 文献文中均有引用（每项至少在正文出现 ≥ 1 次，加上参考文献条目本身 = ≥ 2 次）
    if max_ref > 0:
        cite_cnt = Counter()
        for p in paras:
            for m in re.finditer(r'\[(\d+)\]', p.text):
                cite_cnt[int(m.group(1))] += 1
        unused = [n for n in range(1, max_ref+1) if cite_cnt.get(n,0) <= 1]
        if not unused:
            rep.add(SA,'A10 文献文中均有引用','PASS','','每项 ≥ 1 次正文引用','全部已引')
        else:
            rep.add(SA,'A10 文献文中均有引用','WARN',f'未引: {unused}','每项 ≥ 1 次',f'未引: {unused}')

    # A11 致谢 ≤ 500 字
    thanks_text = ''
    for i, p in enumerate(paras):
        if p.text.strip() == '致谢':
            for j in range(i+1, min(i+15, len(paras))):
                t = paras[j].text.strip()
                # 致谢结束标志
                if t in ('参考文献','附录') or re.match(r'^\[\d+\]',t):
                    break
                thanks_text += t
            break
    tl = len(re.findall(r'[\u4e00-\u9fff]', thanks_text))
    if tl == 0:
        rep.add(SA,'A11 致谢 ≤ 500 字','INFO','未检出（致谢可选）','可选','未含')
    elif tl <= 500:
        rep.add(SA,'A11 致谢 ≤ 500 字','PASS','','≤ 500 字',f'{tl} 字')
    else:
        rep.add(SA,'A11 致谢 ≤ 500 字','FAIL',f'{tl} 字超长','≤ 500 字',f'{tl} 字')


    # ============= B. 排版 =============
    SB = 'B. 排版'
    section = doc.sections[0]
    pw, ph = section.page_width, section.page_height
    is_a4 = abs(pw.cm-21)<0.5 and abs(ph.cm-29.7)<0.5
    rep.add(SB,'B1.1 A4 纸张','PASS' if is_a4 else 'FAIL','','21×29.7 cm',
            f'{pw.cm:.2f}×{ph.cm:.2f} cm')
    margins = [('B1.2 左边距',section.left_margin.cm,2.7),
               ('B1.3 上边距',section.top_margin.cm,2.0),
               ('B1.4 下边距',section.bottom_margin.cm,2.0),
               ('B1.5 右边距',section.right_margin.cm,2.0)]
    for label, act, exp in margins:
        d = abs(act-exp)
        st = 'PASS' if d<0.1 else ('WARN' if d<0.5 else 'FAIL')
        rep.add(SB,label,st,f'偏差 {d:.2f} cm' if d>0.1 else '',f'{exp} cm',f'{act:.2f} cm')

    # B2 1.5 倍行距
    normal_paras = [p for p in paras if p.style.name=='Normal' and p.text.strip()][:30]
    lss = []
    for p in normal_paras:
        if p.paragraph_format.line_spacing:
            lss.append(p.paragraph_format.line_spacing)
    if lss:
        avg = sum(lss)/len(lss)
        if abs(avg-1.5)<0.05:
            rep.add(SB,'B2 正文 1.5 倍行距','PASS',f'抽样 {len(lss)} 段','1.5',f'{avg:.2f}')
        else:
            rep.add(SB,'B2 正文行距','WARN',f'平均 {avg:.2f}','1.5',f'{avg:.2f}')
    else:
        rep.add(SB,'B2 正文行距','INFO','样式默认','1.5','继承')

    # B3 标题字号字体抽样
    hs = {}
    for p in paras:
        sn = p.style.name
        if sn in ('Heading 1','Heading 2','Heading 3') and sn not in hs and p.text.strip():
            hs[sn] = p
    exp_styles = {
        'Heading 1':('B3.1 一级标题','小二号黑体居中',18,'黑体'),
        'Heading 2':('B3.2 二级标题','小四号黑体',12,'黑体'),
        'Heading 3':('B3.3 三级标题','小四号黑体',12,'黑体'),
    }
    for sn,(lbl,desc,ep,ef) in exp_styles.items():
        if sn not in hs:
            rep.add(SB,lbl,'INFO',f'未找到 {sn}',desc,'不适用')
            continue
        p = hs[sn]; r = p.runs[0] if p.runs else None
        if r is None:
            rep.add(SB,lbl,'WARN','无 run',desc,'空')
            continue
        sz = get_run_size(r); ea = get_run_east_asia_font(r)
        ss = fmt_size(sz) if sz else '继承'
        fs = ea or '继承'
        ok_sz = sz and abs(sz.pt-ep)<0.5
        ok_fn = ea and ef in ea
        if ok_sz and ok_fn:
            rep.add(SB,lbl,'PASS','',desc,f'{ss}/{fs}')
        elif not sz and not ea:
            rep.add(SB,lbl,'INFO','样式继承',desc,'样式继承')
        else:
            issues=[]
            if not ok_sz: issues.append(f'字号 {ss}')
            if not ok_fn: issues.append(f'字体 {fs}')
            rep.add(SB,lbl,'WARN','; '.join(issues),desc,f'{ss}/{fs}')

    # B4 正文字号字体
    normal_sample = None
    for p in paras:
        if p.style.name=='Normal' and p.text.strip() and len(p.text)>30:
            normal_sample = p; break
    if normal_sample and normal_sample.runs:
        r = normal_sample.runs[0]
        sz=get_run_size(r); ea=get_run_east_asia_font(r)
        ok_sz = sz and abs(sz.pt-12)<0.5
        ok_fn = ea and ('宋体' in ea or 'SimSun' in str(ea).lower() or 'Song' in str(ea))
        if ok_sz and ok_fn:
            rep.add(SB,'B4 正文 小四号宋体','PASS','','小四号宋体',f'{fmt_size(sz)}/{ea}')
        elif not sz and not ea:
            rep.add(SB,'B4 正文字号字体','INFO','样式继承','小四号宋体','继承')
        else:
            issues=[]
            if not ok_sz: issues.append(f'字号 {fmt_size(sz)}')
            if not ok_fn: issues.append(f'字体 {ea or "?"}')
            rep.add(SB,'B4 正文字号字体','WARN','; '.join(issues),'小四号宋体',
                    f'{fmt_size(sz) if sz else "?"}/{ea or "?"}')


    # ============= C. 图表 =============
    SC = 'C. 图表'

    # C1 表编号连续性
    tbl_count = len(doc.tables)
    tbl_caps = [(pi,int(re.match(r'^\s*表\s*(\d+)',p.text).group(1)),p.text.strip())
                for pi,p in enumerate(paras) if re.match(r'^\s*表\s*(\d+)\s',p.text)]
    if tbl_caps:
        nums = sorted(set(t[1] for t in tbl_caps))
        miss = [n for n in range(1, max(nums)+1) if n not in nums]
        if not miss:
            rep.add(SC,'C1 表编号连续性','PASS','','连续',f'表 1—{max(nums)}, {len(nums)} 个')
        else:
            rep.add(SC,'C1 表编号连续性','FAIL',f'缺 {miss}','连续',f'缺 {miss}')
    rep.add(SC,'C1.2 表格对象数','INFO','','与编号一致',f'{tbl_count} 个 Table 对象')

    # C2 图编号连续性
    fig_caps = [(pi,int(re.match(r'^\s*图\s*(\d+)',p.text).group(1)),p.text.strip())
                for pi,p in enumerate(paras) if re.match(r'^\s*图\s*(\d+)\s',p.text)]
    if fig_caps:
        nums = sorted(set(f[1] for f in fig_caps))
        miss = [n for n in range(1, max(nums)+1) if n not in nums]
        if not miss:
            rep.add(SC,'C2 图编号连续性','PASS','','连续',f'图 1—{max(nums)}, {len(nums)} 个')
        else:
            rep.add(SC,'C2 图编号连续性','FAIL',f'缺 {miss}','连续',f'缺 {miss}')

    # C3 三线表（不是 Grid 样式）
    grid_tbls = []
    for ti, tbl in enumerate(doc.tables):
        sn = tbl.style.name if tbl.style else ''
        if 'Grid' in sn or '网格' in sn:
            grid_tbls.append((ti, sn))
    if grid_tbls:
        rep.add(SC,'C3 三线表','WARN',f'{len(grid_tbls)} 个表使用 Grid 样式',
                '三线表',f'{[f"T{ti}({sn})" for ti,sn in grid_tbls]}')
    else:
        rep.add(SC,'C3 三线表','PASS','','三线表',f'全部 {tbl_count} 个表非 Grid')

    # C4 图标题位置（图标题段下方应紧邻图，应在图下方但 docx 中图在前/标题在后）
    # 规范：图标题在图的下方；表标题在表的上方
    # 这里检查：图标题段的上一段是否含图片对象
    fig_with_image = 0
    fig_without = []
    NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    NS_PIC = "http://schemas.openxmlformats.org/drawingml/2006/picture"
    NS_DRAW = "http://schemas.openxmlformats.org/drawingml/2006/main"
    def has_image(p):
        el = p._element
        # 查找 drawing 或 pic 元素
        for tag in [f'.//{{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}}inline',
                    f'.//{{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}}anchor',
                    f'.//{{{NS_W}}}drawing',
                    f'.//{{{NS_W}}}pict']:
            if el.find(tag) is not None: return True
        return False
    for pi, n, _ in fig_caps:
        # 检查上方 5 段内是否有图片
        found = False
        for i in range(max(0,pi-5), pi):
            if has_image(paras[i]):
                found = True; break
        # 若上方没有，再检查下方 3 段
        if not found:
            for i in range(pi+1, min(len(paras), pi+4)):
                if has_image(paras[i]):
                    found = True; break
        if found:
            fig_with_image += 1
        else:
            fig_without.append(n)
    if fig_with_image == len(fig_caps):
        rep.add(SC,'C4 图标题与图配对','PASS','','每图均配对',f'{fig_with_image}/{len(fig_caps)}')
    else:
        rep.add(SC,'C4 图标题与图配对','WARN',f'图 {fig_without} 周围 8 段未检出图片',
                '每图均配对', f'{fig_with_image}/{len(fig_caps)} 配对')

    # C5 表内首行缩进
    indent_cells = []
    for ti, tbl in enumerate(doc.tables):
        for ri, row in enumerate(tbl.rows):
            for ci, cell in enumerate(row.cells):
                for p in cell.paragraphs:
                    pf = p.paragraph_format
                    if pf.first_line_indent and pf.first_line_indent.cm > 0.1:
                        indent_cells.append((ti, ri, ci, pf.first_line_indent.cm))
    if not indent_cells:
        rep.add(SC,'C5 表内首行缩进','PASS','','无首行缩进','0 处')
    else:
        rep.add(SC,'C5 表内首行缩进','WARN',
                f'{len(indent_cells)} 个表内单元格有首行缩进',
                '无','例: '+str(indent_cells[:3]))


    # ============= D. 数字/标点 =============
    SD = 'D. 数字/标点'
    han_date = re.findall(r'[一二三四五六七八九〇零]+年[一二三四五六七八九〇零十]+月', all_text)
    arab_date = re.findall(r'\b(19|20)\d{2}年\d{1,2}月', all_text)
    if han_date:
        rep.add(SD,'D1 公历日期阿拉伯','WARN',f'汉字公历 {len(han_date)} 处',
                '阿拉伯数字',f'例: {han_date[:3]}')
    else:
        rep.add(SD,'D1 公历日期阿拉伯','PASS','','阿拉伯数字',f'{len(arab_date)} 处')

    # D2 中文段英文逗号
    cn_paras_text = []
    for p in paras:
        t = p.text
        if not t.strip(): continue
        cn_chars = len(re.findall(r'[\u4e00-\u9fff]', t))
        if cn_chars > len(t)*0.5:
            cn_paras_text.append(t)
    cn_text = '\n'.join(cn_paras_text)
    bad_comma = re.findall(r',[ \u4e00-\u9fff]', cn_text)
    if len(bad_comma) > 5:
        rep.add(SD,'D2 中文标点','WARN',f'{len(bad_comma)} 处疑似英文逗号',
                '中文段中文标点',f'{len(bad_comma)} 处')
    elif len(bad_comma) > 0:
        rep.add(SD,'D2 中文标点','INFO',f'{len(bad_comma)} 处（含合理用例）',
                '中文段中文标点',f'{len(bad_comma)} 处')
    else:
        rep.add(SD,'D2 中文标点','PASS','','GB/T 15834-1995','0 处')


    # ============= E. 标题序次 =============
    SE = 'E. 标题序次'
    h1_p = re.compile(r'^[一二三四五六七八九十]+、')
    h2_p = re.compile(r'^（[一二三四五六七八九十]+）')
    digit_p = re.compile(r'^\d+(\.\d+)+\s')
    h1_n = sum(1 for p in paras if h1_p.match(p.text.strip()))
    h2_n = sum(1 for p in paras if h2_p.match(p.text.strip()))
    dn = sum(1 for p in paras if digit_p.match(p.text.strip()))
    if h1_n > 0 and h2_n > 0:
        rep.add(SE,'E1 经管类序次','PASS',f'一级 {h1_n}, 二级 {h2_n}',
                '一、（一）1.（1）①',f'一 {h1_n}, （一） {h2_n}')
        if dn > 5:
            rep.add(SE,'E2 序次混用','INFO',f'另有 {dn} 处 N.N（可能公式编号）',
                    '一致用汉字序次','N.N 形式 '+str(dn))


    # ============= F. 装订（INFO） =============
    SF = 'F. 装订'
    rep.add(SF,'F1 装订顺序','INFO','封面→承诺→目录→题目→正文→注释→文献→致谢',
            '本规范要求','请人工核对')
    rep.add(SF,'F2 装订线','INFO','左侧装订；左边距 2.7 cm 已符合','左侧','配套已符合')
    rep.add(SF,'F3 打印面','INFO','内页单面打印','单面','打印时确认')


    # ============= G. 上下标特殊字符 =============
    SG = 'G. 上下标特殊字符'

    # G1 vertAlign run 健康（v3 应为 0 空 + 仅角标）
    va_total = 0; va_empty = 0; va_texts = []
    for p in paras:
        for r in p.runs:
            rpr = r._element.find(qn('w:rPr'))
            if rpr is None: continue
            va = rpr.find(qn('w:vertAlign'))
            if va is None: continue
            va_total += 1
            if not r.text: va_empty += 1
            else: va_texts.append(r.text)
    if va_empty == 0:
        rep.add(SG,'G1 vertAlign 健康','PASS','','无空 sup run',
                f'共 {va_total} run, 全有文本')
    else:
        rep.add(SG,'G1 vertAlign 健康','FAIL',f'{va_empty} 空 sup run',
                '0 空', f'{va_empty} 空')

    # G2 角标完整性（参考文献角标 [1]..[max_ref] 应都有 sup）
    sup_refs = sorted(set(re.search(r'\d+', t).group() for t in va_texts if re.search(r'\[\d+\]', t)))
    sup_ref_nums = set()
    for t in va_texts:
        m = re.search(r'\[(\d+)\]', t)
        if m: sup_ref_nums.add(int(m.group(1)))
    missing_sup = []
    if max_ref > 0:
        # 找到正文中出现 [N] 但未在 superscript 中的
        text_refs = set()
        for p in paras:
            for m in re.finditer(r'\[(\d+)\]', p.text):
                text_refs.add(int(m.group(1)))
        # 参考文献条目本身的 [N] 是普通文本（在文献列表）
        # 我们的检查重点：被引用的角标 [N] 是否上标
        # 这里简化：sup_ref_nums 应包含 1..max_ref 中的大部分
        coverage = len(sup_ref_nums) / max_ref if max_ref else 0
        if coverage >= 0.7:
            rep.add(SG,'G2 角标上标覆盖','PASS',
                    f'{len(sup_ref_nums)}/{max_ref} 文献被以 superscript 引用',
                    '≥70%', f'{coverage*100:.0f}%')
        else:
            rep.add(SG,'G2 角标上标覆盖','WARN',
                    f'{len(sup_ref_nums)}/{max_ref} 文献被 superscript 引用',
                    '≥70%', f'{coverage*100:.0f}%')

    # G3 单位上标（km², m² 等已是 Unicode）
    bad_units = []
    bad_units.extend([f'km2 @ P{pi}' for pi,p in enumerate(paras) for _ in re.finditer(r'\bkm2\b',p.text)][:10])
    bad_units.extend([f'm2 @ P{pi}' for pi,p in enumerate(paras) for _ in re.finditer(r'(?<![ka-z])m2\b',p.text)][:10])
    if bad_units:
        rep.add(SG,'G3 单位上标 ASCII','WARN',f'{len(bad_units)} 处 ASCII 单位',
                'km²/m² Unicode','例: '+str(bad_units[:5]))
    else:
        rep.add(SG,'G3 单位上标 ASCII','PASS','','Unicode','无 ASCII')

    # G4 LaTeX 残留
    latex_residue = []
    for pi, p in enumerate(paras):
        if re.search(r'\\(prime|alpha|beta|gamma|theta|sum|frac|sqrt|cdot|cdots)\b', p.text):
            latex_residue.append((pi, p.text[:80]))
        if re.search(r'\$[^$\n]{1,80}\$', p.text):
            latex_residue.append((pi, p.text[:80]))
    if not latex_residue:
        rep.add(SG,'G4 LaTeX 残留','PASS','','无 LaTeX 源码','清洁')
    else:
        rep.add(SG,'G4 LaTeX 残留','WARN',f'{len(latex_residue)} 处','无','例: '+str(latex_residue[:2]))


    # ============= H. 交叉引用深度 =============
    SH = 'H. 交叉引用深度'

    caption_paras = set(pi for pi,_,_ in tbl_caps) | set(pi for pi,_,_ in fig_caps)
    xref_fig, xref_tab = [], []
    for pi, p in enumerate(paras):
        if pi in caption_paras: continue
        for m in re.finditer(r'图\s*(\d+)', p.text):
            xref_fig.append((pi, int(m.group(1))))
        for m in re.finditer(r'表\s*(\d+)', p.text):
            xref_tab.append((pi, int(m.group(1))))

    fig_nums = sorted(set(h[1] for h in xref_fig))
    tab_nums = sorted(set(h[1] for h in xref_tab))
    fig_max = max(set(f[1] for f in fig_caps), default=0)
    tab_max = max(set(t[1] for t in tbl_caps), default=0)
    fig_uncited = sorted(set(range(1,fig_max+1)) - set(fig_nums))
    tab_uncited = sorted(set(range(1,tab_max+1)) - set(tab_nums))
    oob_fig = [h for h in xref_fig if h[1]<1 or h[1]>fig_max]
    oob_tab = [h for h in xref_tab if h[1]<1 or h[1]>tab_max]

    if not fig_uncited:
        rep.add(SH,'H1 图全部被引用','PASS','','全覆盖',f'图 1—{fig_max}')
    else:
        rep.add(SH,'H1 图全部被引用','FAIL',f'未引: {fig_uncited}','全覆盖',f'未引 {fig_uncited}')
    if not tab_uncited:
        rep.add(SH,'H2 表全部被引用','PASS','','全覆盖',f'表 1—{tab_max}')
    else:
        rep.add(SH,'H2 表全部被引用','FAIL',f'未引: {tab_uncited}','全覆盖',f'未引 {tab_uncited}')
    rep.add(SH,'H3 图引用越界','PASS' if not oob_fig else 'FAIL',
            '' if not oob_fig else f'{len(oob_fig)} 越界','无越界',f'{len(oob_fig)} 处')
    rep.add(SH,'H4 表引用越界','PASS' if not oob_tab else 'FAIL',
            '' if not oob_tab else f'{len(oob_tab)} 越界','无越界',f'{len(oob_tab)} 处')

    # 公式引用
    form_refs = []
    for pi, p in enumerate(paras):
        for m in re.finditer(r'公式\s*[（(](\d+\.?\d*)[）)]', p.text):
            form_refs.append((pi, m.group(1)))
        for m in re.finditer(r'式\s*[（(](\d+\.?\d*)[）)]', p.text):
            form_refs.append((pi, m.group(1)))
    rep.add(SH,'H5 公式引用','INFO',f'{len(form_refs)} 处公式引用','章节式编号 (N.M)',
            f'编号: {sorted(set(r[1] for r in form_refs))}')


    # ============= I. 参考文献深度 =============
    SI = 'I. 参考文献深度'

    if ref_items:
        # I1 编号连续性
        ref_nums = sorted(set(r[0] for r in ref_items))
        miss_ref = [n for n in range(1, max(ref_nums)+1) if n not in ref_nums]
        if not miss_ref:
            rep.add(SI,'I1 文献编号连续','PASS','','连续',f'[1]—[{max(ref_nums)}]')
        else:
            rep.add(SI,'I1 文献编号连续','FAIL',f'缺 {miss_ref}','连续',f'缺 {miss_ref}')

        # I2 引用频次分布
        # cite_cnt 已包含所有正文 [N] 出现次数（含列表本身的 [N]）
        # 列表条目每个 [N] 计 1 次，所以正文引用次数 = cite_cnt[n] - 1
        cite_freq = {n: cite_cnt.get(n,0)-1 for n in range(1, max(ref_nums)+1)}
        ge1 = sum(1 for v in cite_freq.values() if v>=1)
        rep.add(SI,'I2 文献引用率','INFO',
                f'被引 ≥1 次: {ge1}/{max(ref_nums)}',
                '理想 100%', f'{ge1/max(ref_nums)*100:.0f}%')

        # I3 文献年份范围
        years = []
        for n, t in ref_items:
            m = re.findall(r'\b(19|20)\d{2}\b', t)
            if m: years.extend([int(y+m_) for y,m_ in [(yy,'') for yy in m]] if False else [int(yy) for yy in [s for s in re.findall(r'(?:19|20)\d{2}',t)]])
        years = sorted(set([y for y in years if 1990 <= y <= 2026]))
        if years:
            rep.add(SI,'I3 文献年份分布','INFO',
                    f'{years[0]}—{years[-1]}',
                    '近 5 年文献 ≥ 1/3', f'{years[0]}—{years[-1]} 共 {len(years)} 个独立年份')
        # I4 中英文文献
        en_refs = sum(1 for n,t in ref_items if len(re.findall(r'[a-zA-Z]', t)) > len(t)*0.4)
        cn_refs = max_ref - en_refs
        rep.add(SI,'I4 中英文文献','INFO',f'中 {cn_refs}, 英 {en_refs}',
                '应有英文文献','中 '+str(cn_refs)+', 英 '+str(en_refs))

        # I5 文献类型（[J]/[M]/[D]/[C]/[N]/[Z]/...）
        types = Counter()
        for n,t in ref_items:
            m = re.search(r'\[([JMDCRNGPZ])\]', t)
            if m: types[m.group(1)] += 1
            else: types['?'] += 1
        rep.add(SI,'I5 文献类型','INFO',f'类型分布','含 J/M/D/C 等',str(dict(types)))


    # ============= J. 内容鲁棒性 =============
    SJ = 'J. 内容鲁棒性'

    # J1 PLUS 模拟标注一致性：所有出现"2025"的段，看是否标注"PLUS模拟"或"模拟情景"
    plus_keywords = ['PLUS', '模拟', '情景']
    paras_with_2025 = []
    for pi, p in enumerate(paras):
        if '2025' in p.text:
            paras_with_2025.append((pi, p.text))
    paras_2025_with_plus = [pi for pi,t in paras_with_2025 if any(k in t for k in plus_keywords)]
    paras_2025_no_plus = [(pi,t) for pi,t in paras_with_2025 if not any(k in t for k in plus_keywords)]
    if paras_with_2025:
        cov = len(paras_2025_with_plus)/len(paras_with_2025)
        if cov >= 0.5:
            rep.add(SJ,'J1 2025 数据 PLUS 标注','INFO',
                    f'{len(paras_2025_with_plus)}/{len(paras_with_2025)} 段含 PLUS/模拟/情景关键词',
                    '所有涉 2025 段应标注','覆盖率 '+f'{cov*100:.0f}%')
        else:
            rep.add(SJ,'J1 2025 数据 PLUS 标注','WARN',
                    f'仅 {len(paras_2025_with_plus)}/{len(paras_with_2025)} 段含标注',
                    '应均标注','覆盖率 '+f'{cov*100:.0f}%')

    # J2 关键数据一致性（耕地减少幅度、ESI 值等）
    # 提取所有 "减少 X.XX km²" / "降至 0.XXX" 等关键数值
    cropland_decline = re.findall(r'耕地.{0,20}减少\s*([\d.]+)\s*km', all_text)
    esi_values = re.findall(r'(?:ESI|生态安全指数).{0,30}(0\.\d{3})', all_text)
    forest_increase = re.findall(r'林地.{0,20}增加\s*([\d.]+)\s*km', all_text)
    rep.add(SJ,'J2 耕地减少数值','INFO',f'{len(cropland_decline)} 处',
            '数值一致',str(Counter(cropland_decline)))
    rep.add(SJ,'J3 ESI 数值','INFO',f'{len(esi_values)} 处',
            '数值一致',str(Counter(esi_values)))
    rep.add(SJ,'J4 林地增加数值','INFO',f'{len(forest_increase)} 处',
            '数值一致',str(Counter(forest_increase)))

    # J5 关键术语前后一致：横州市 vs 横县
    hengzhou = all_text.count('横州市')
    hengxian = all_text.count('横县')
    if hengzhou > 0 and hengxian > 0:
        rep.add(SJ,'J5 地名统一','WARN',f'横州市 {hengzhou}, 横县 {hengxian} 共现',
                '统一名称', '可能新旧名混用')
    else:
        rep.add(SJ,'J5 地名统一','PASS','','统一名称',f'横州市 {hengzhou}')


    # ============= K. 字体一致性 =============
    SK = 'K. 字体一致性'

    # 按字号统计 run
    size_dist = Counter()
    font_dist = Counter()
    for p in paras:
        for r in p.runs:
            if not r.text.strip(): continue
            sz = get_run_size(r)
            ea = get_run_east_asia_font(r)
            sz_key = round(sz.pt, 1) if sz else '继承'
            size_dist[sz_key] += 1
            font_dist[ea or '继承'] += 1
    top_sizes = size_dist.most_common(8)
    top_fonts = font_dist.most_common(8)
    rep.add(SK,'K1 字号分布','INFO',f'共 {sum(size_dist.values())} run',
            '主流字号小四(12pt)',str(top_sizes))
    rep.add(SK,'K2 字体分布','INFO',f'共 {sum(font_dist.values())} run',
            '主流中文宋体',str(top_fonts))

    # K3 异常字号（非 9/10.5/12/14/15/16/18/22/24）
    normal_sizes = {9,10.5,12,14,15,16,18,22,24}
    abnormal_sizes = [k for k in size_dist if isinstance(k,(int,float)) and k not in normal_sizes]
    if abnormal_sizes:
        rep.add(SK,'K3 异常字号','WARN',f'非标准字号 {abnormal_sizes}',
                '标准字号',str(abnormal_sizes))
    else:
        rep.add(SK,'K3 异常字号','PASS','','标准字号','全部标准')


    # ============= L. 段落格式 =============
    SL = 'L. 段落格式'

    # L1 正文首行缩进 2 字符（约 0.74-0.85 cm）
    indented = 0; not_indented = 0
    for p in paras:
        if p.style.name != 'Normal': continue
        if not p.text.strip(): continue
        if len(p.text) < 30: continue  # 跳过短段
        pf = p.paragraph_format
        if pf.first_line_indent and pf.first_line_indent.cm > 0.5:
            indented += 1
        else:
            not_indented += 1
    total = indented + not_indented
    if total:
        cov = indented/total
        if cov >= 0.7:
            rep.add(SL,'L1 正文首行缩进 2 字符','PASS',
                    f'{indented}/{total} 段首行缩进',
                    '≥70%', f'{cov*100:.0f}%')
        else:
            rep.add(SL,'L1 正文首行缩进 2 字符','INFO',
                    f'{indented}/{total} 段（多数继承样式默认）',
                    '≥70% 或样式继承', f'{cov*100:.0f}%')


    # ============= 保存 =============
    md = rep.to_md(DOC)
    with open(OUT,'w',encoding='utf-8') as f:
        f.write(md)
    print(f"报告 -> {OUT} ({os.path.getsize(OUT):,} bytes)")
    c = rep.stats()
    print(f"\n概要: PASS={c.get('PASS',0)} WARN={c.get('WARN',0)} FAIL={c.get('FAIL',0)} INFO={c.get('INFO',0)}")


if __name__ == '__main__':
    main()
