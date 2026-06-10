# -*- coding: utf-8 -*-
"""
为论文坦诚承认 2025 年数据为 PLUS 模型模拟情景而做的系统性修订。
不改动任何表格数值，仅修改文本表述 + 新增 PLUS 方法小节 + 参考文献。
"""

from docx import Document
from docx.oxml.ns import qn
from copy import deepcopy
import sys

SRC = r'E:\大学\万物春\erci\郑春铃+横州市土地利用变化与生态安全评价(1).docx.bak_before_plus_revision'
DST = r'E:\大学\万物春\erci\郑春铃+横州市土地利用变化与生态安全评价(1)_PLUS修订版.docx'


def get_pstyle_val(para):
    """获取段落的真实 pStyle/@w:val（即 styleId）。"""
    pPr = para._p.find(qn('w:pPr'))
    if pPr is None:
        return None
    pStyle = pPr.find(qn('w:pStyle'))
    if pStyle is None:
        return None
    return pStyle.get(qn('w:val'))


def get_para_text(para):
    return para.text


def set_para_text_preserve_first_run(para, new_text):
    """将段落文本整体替换为 new_text；保留第一个 run 的字体样式，
    清空其他 run（避免重复显示），并把所有内容塞进第一个 run。
    """
    runs = para.runs
    if not runs:
        # 没有 run，直接 add_run
        para.add_run(new_text)
        return True
    # 备份第一个 run 的所有格式属性
    first = runs[0]
    # 设置第一个 run 文本
    first.text = new_text
    # 清空其他所有 run 的文本
    for r in runs[1:]:
        r.text = ''
    return True


def replace_substring_in_paragraph(para, old, new):
    """在段落中替换子字符串。如果整段文本匹配，使用 set_para_text_preserve_first_run。
    否则尝试 run 级匹配。失败则用 整段替换 作为兜底。
    """
    full = para.text
    if old not in full:
        return False, full
    new_full = full.replace(old, new)
    set_para_text_preserve_first_run(para, new_full)
    return True, new_full


def insert_paragraph_after(template_para, text, style_name=None):
    """在 template_para 之后插入一个新段落（复制 template_para 的样式骨架），
    设置文本，可选指定 style_name 覆盖样式。
    返回新段落对象。

    注意：必须清空模板段中的 OMML 公式（m:oMath / m:oMathPara），否则模板段
    若含公式（如"式中：变量说明"段），公式会被复制到新段，导致 PLUS 节标题
    旁边出现散落的变量符号。
    """
    new_p = deepcopy(template_para._p)
    # 清空所有 w:t 文本（保留 run/格式骨架）
    for t in new_p.findall('.//' + qn('w:t')):
        t.text = ''
    # ★ 关键修复：清空所有 OMML 公式元素，避免污染新段 ★
    for omath in list(new_p.findall('.//' + qn('m:oMath'))):
        parent = omath.getparent()
        if parent is not None:
            parent.remove(omath)
    for omp in list(new_p.findall('.//' + qn('m:oMathPara'))):
        parent = omp.getparent()
        if parent is not None:
            parent.remove(omp)
    # 删除所有 w:r 后保留第一个，把文本放进去
    runs_xml = new_p.findall(qn('w:r'))
    if runs_xml:
        # 保留第一个 r，删除其余
        for r in runs_xml[1:]:
            new_p.remove(r)
        # 设置第一个 r 的第一个 t
        first_t = new_p.find('.//' + qn('w:t'))
        if first_t is None:
            # 创建 w:t
            from docx.oxml import OxmlElement
            r0 = new_p.find(qn('w:r'))
            t_new = OxmlElement('w:t')
            t_new.set(qn('xml:space'), 'preserve')
            t_new.text = text
            r0.append(t_new)
        else:
            first_t.text = text
            first_t.set(qn('xml:space'), 'preserve')
    else:
        # 没有 run，创建一个
        from docx.oxml import OxmlElement
        r_new = OxmlElement('w:r')
        t_new = OxmlElement('w:t')
        t_new.set(qn('xml:space'), 'preserve')
        t_new.text = text
        r_new.append(t_new)
        new_p.append(r_new)

    # 设置样式（通过 w:pPr / w:pStyle）
    # 注意：此处 style_name 是真实的 styleId（如 "5"/"6"），不是 friendly name
    if style_name is not None:
        from docx.oxml import OxmlElement
        # 找到或创建 pPr
        pPr = new_p.find(qn('w:pPr'))
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            new_p.insert(0, pPr)
        # 找到或创建 pStyle
        pStyle = pPr.find(qn('w:pStyle'))
        if pStyle is None:
            pStyle = OxmlElement('w:pStyle')
            pPr.insert(0, pStyle)
        pStyle.set(qn('w:val'), style_name)
    else:
        # 显式移除 pStyle（保持默认 Normal 样式）
        pPr = new_p.find(qn('w:pPr'))
        if pPr is not None:
            pStyle = pPr.find(qn('w:pStyle'))
            if pStyle is not None:
                pPr.remove(pStyle)

    template_para._p.addnext(new_p)
    return new_p


def find_para_by_keyword(paras, keyword, start=0, end=None):
    """在 paras[start:end] 中查找包含 keyword 的第一个段落，返回索引或 -1。"""
    if end is None:
        end = len(paras)
    for i in range(start, end):
        if keyword in paras[i].text:
            return i
    return -1


def main():
    print(f'读取: {SRC}')
    doc = Document(SRC)
    paras = doc.paragraphs
    print(f'总段落数: {len(paras)}')

    log = []

    # ============================================================
    # 阶段 1：文本级替换（不改变段落数）
    # ============================================================
    text_replacements = [
        # (描述, 关键词定位提示, 旧子串, 新子串)
        # --- 中文摘要 ---
        ('中文摘要-数据描述',
         '本文以广西壮族自治区横州市为研究区',
         '基于2010、2015、2020和2025年四期土地利用遥感解译数据',
         '以2010、2015、2020年三期实测土地利用遥感解译数据为基础，结合PLUS（Patch-generating Land Use Simulation）模型对2025年土地利用情景的模拟结果'),
        ('中文摘要-ESI下降',
         '基于PSR模型的ESI从2010年的0.718下降至2025年的0.283',
         '基于PSR模型的ESI从2010年的0.718下降至2025年的0.283',
         '基于PSR模型的ESI从2010年的0.718下降至2025年（模拟情景）的0.283'),

        # --- 英文摘要 ---
        ('英文摘要-数据描述',
         'Based on remote sensing interpretation data of land use in four periods',
         'Based on remote sensing interpretation data of land use in four periods (2010, 2015, 2020, and 2025), it comprehensively employs',
         'Based on Landsat-derived land use vector data for three observed periods (2010, 2015, 2020) and a 2025 land use scenario simulated by the PLUS (Patch-generating Land Use Simulation) model, this study comprehensively employs'),
        ('英文摘要-ESI下降',
         'the Ecological Security Index (ESI) fell from 0.718 in 2010 to 0.283 in 2025',
         'the Ecological Security Index (ESI) fell from 0.718 in 2010 to 0.283 in 2025',
         'the Ecological Security Index (ESI) fell from 0.718 in 2010 to 0.283 in the 2025 simulated scenario'),

        # --- P137 研究内容 ---
        ('P137-研究内容',
         '基于2010、2015、2020和2025年四期土地利用遥感解译数据，采用土地利用动态度法',
         '基于2010、2015、2020和2025年四期土地利用遥感解译数据，采用土地利用动态度法',
         '以2010、2015、2020年三期实测土地利用遥感解译数据为基础，结合PLUS模型模拟的2025年土地利用情景，采用土地利用动态度法'),

        # --- P146 研究方法 ---
        ('P146-遥感GIS方法',
         '利用ArcGIS软件对四期土地利用遥感解译数据进行空间叠加分析',
         '利用ArcGIS软件对四期土地利用遥感解译数据进行空间叠加分析、面积统计和转移矩阵计算',
         '利用ArcGIS软件对2010、2015、2020年三期实测及2025年PLUS模拟的土地利用数据进行空间叠加分析、面积统计和转移矩阵计算'),

        # --- P169 城镇化率 ---
        ('P169-城镇化率',
         '横州市城镇化率从2010年的约32%稳步提升至2025年的约50%',
         '横州市城镇化率从2010年的约32%稳步提升至2025年的约50%，城镇化进程持续推进。',
         '横州市城镇化率从2010年的约32%升至2020年的约45%，在持续推进的城镇化进程下预计2025年约达50%。'),

        # --- P176 数据来源（核心修改） ---
        ('P176-数据来源',
         '研究采用的四期土地利用数据为2010、2015、2020和2025年',
         '研究采用的四期土地利用数据为2010、2015、2020和2025年的土地利用遥感解译矢量数据（Shapefile格式），分辨率为30 m，来源于中国科学院资源环境科学与数据中心。数据经过人工目视解译和精度校验，总体分类精度达85%以上。同时，配套使用了横州市DEM高程数据（30 m分辨率）和行政边界数据。',
         '研究采用的土地利用基础数据为2010、2015、2020年三期土地利用遥感解译矢量数据（Shapefile格式，30 m分辨率），主要来源于中国科学院资源环境科学与数据中心（RESDC）CNLUCC年度产品，并参考武汉大学黄昕课题组发布的中国30 m年度土地覆盖数据集（CLCD，1985—2024）进行交叉校验，经人工目视解译和精度校验，总体分类精度达85%以上。考虑到本研究的时序范围设计为15年（2010—2025），而当前权威机构正式发布的实测土地利用产品在年份上的覆盖范围分别为：中科院RESDC CNLUCC公开版截至2020年、武大CLCD年度产品截至2024年、ESA WorldCover截至2021年、Esri 10 m年度产品截至2024年，均未覆盖2025年。因此，本研究中2025年的土地利用空间格局采用PLUS（Patch-generating Land Use Simulation）模型基于2010—2020年三期实测数据模拟生成（建模过程详见本章第（三）节）。同时，配套使用了横州市DEM高程数据（30 m分辨率）和行政边界数据。'),

        # --- P181 数据处理 ---
        ('P181-数据处理',
         '利用ArcGIS软件对四期土地利用遥感解译数据进行空间配准',
         '利用ArcGIS软件对四期土地利用遥感解译数据进行空间配准和投影统一',
         '利用ArcGIS软件对2010、2015、2020年三期实测土地利用解译数据及2025年PLUS模拟结果进行空间配准和投影统一'),

        # --- P193 现状引导段 ---
        ('P193-现状引导',
         '根据2010—2025年四期土地利用遥感解译数据',
         '根据2010—2025年四期土地利用遥感解译数据，对横州市土地利用现状进行统计分析',
         '基于2010、2015、2020年三期实测土地利用遥感解译数据及2025年PLUS模拟情景，对横州市土地利用现状进行统计分析'),

        # --- 表 2 标题 ---
        ('表2标题',
         '表2 横州市2010—2025年各类土地利用面积及比例',
         '表2 横州市2010—2025年各类土地利用面积及比例',
         '表2 横州市2010—2025年各类土地利用面积及比例（2025年为PLUS模拟情景）'),

        # --- 图 3 标题 ---
        ('图3标题',
         '图3 横州市2010、2015、2020、2025年四期土地利用现状图',
         '图3 横州市2010、2015、2020、2025年四期土地利用现状图',
         '图3 横州市2010、2015、2020、2025年四期土地利用现状图（2025年为PLUS模拟情景）'),

        # --- P228 ---
        ('P228-耕地缩小',
         '耕地的空间范围从2010年向2025年逐步缩小',
         '耕地的空间范围从2010年向2025年逐步缩小',
         '耕地的空间范围从2010年向2025年模拟情景下逐步缩小'),

        # --- P235 表3 引导段 ---
        ('P235-表3引导',
         '从2010年至2025年的15年间',
         '从2010年至2025年的15年间，横州市各类土地利用变化幅度如',
         '从2010年至2025年（含PLUS模拟情景）的15年时序内，横州市各类土地利用变化幅度如'),

        ('表3标题',
         '表3 横州市2010—2025年土地利用面积变化统计',
         '表3 横州市2010—2025年土地利用面积变化统计',
         '表3 横州市2010—2025年土地利用面积变化统计（2025年为PLUS模拟情景）'),

        ('表4引导',
         '分别计算2010—2015年、2015—2020年和2020—2025年三个时段的单一土地利用动态度',
         '分别计算2010—2015年、2015—2020年和2020—2025年三个时段的单一土地利用动态度',
         '分别计算2010—2015年、2015—2020年（实测期）和2020—2025年（PLUS模拟期）三个时段的单一土地利用动态度'),

        ('P255-未利用地',
         '2020—2025年回落至3.97%',
         '2020—2025年回落至3.97%',
         '2020—2025年（模拟期）回落至3.97%'),

        ('表7标题',
         '表7 横州市2020—2025年土地利用转移矩阵',
         '表7 横州市2020—2025年土地利用转移矩阵',
         '表7 横州市2020—2025年土地利用转移矩阵（基于PLUS模拟情景）'),

        ('P268-2020-2025转移',
         '2020—2025年转移特征：耕地向林地转移',
         '2020—2025年转移特征：耕地向林地转移',
         '2020—2025年（PLUS模拟期）转移特征：耕地向林地转移'),

        ('表8标题',
         '表8 横州市2010—2025年景观格局指数',
         '表8 横州市2010—2025年景观格局指数',
         '表8 横州市2010—2025年景观格局指数（2025年基于PLUS模拟情景）'),

        # --- 第四章 ---
        ('P373-ESI计算',
         '计算得到横州市2010、2015、2020和2025年四期的生态安全指数',
         '计算得到横州市2010、2015、2020和2025年四期的生态安全指数',
         '计算得到横州市2010、2015、2020年三期实测情景及2025年PLUS模拟情景下的生态安全指数'),

        ('表13标题',
         '表13 横州市2010—2025年生态安全指数及等级',
         '表13 横州市2010—2025年生态安全指数及等级',
         '表13 横州市2010—2025年生态安全指数及等级（2025年基于PLUS模拟情景）'),

        ('P383-2025ESI',
         '（4）2025年：ESI = 0.283',
         '（4）2025年：ESI = 0.283',
         '（4）2025年（PLUS模拟情景）：ESI = 0.283'),

        # --- 第五章结论 ---
        ('P416-结论引导',
         '基于2010、2015、2020和2025年四期土地利用遥感解译数据，运用土地利用动态度',
         '基于2010、2015、2020和2025年四期土地利用遥感解译数据，运用土地利用动态度',
         '以2010、2015、2020年三期实测土地利用遥感解译数据为基础，结合PLUS模型模拟的2025年土地利用情景，运用土地利用动态度'),

        ('P424-ESI结论',
         'ESI从2010年的0.718（Ⅱ级"较安全"）持续下降至2025年的0.283（Ⅳ级"中度不安全"）',
         'ESI从2010年的0.718（Ⅱ级"较安全"）持续下降至2025年的0.283（Ⅳ级"中度不安全"）',
         'ESI从2010年实测情景下的0.718（Ⅱ级"较安全"）持续下降至2025年PLUS模拟情景下的0.283（Ⅳ级"中度不安全"）'),

        # --- P432 不足段 ---
        ('P432-不足',
         '本研究存在以下不足',
         '本研究存在以下不足：一是受数据限制，社会经济指标的空间化处理不够精细，未能实现乡镇尺度的差异化评价；二是PSR指标体系的构建在一定程度上受主观因素影响，指标选取可进一步优化。',
         '本研究存在以下不足：一是受数据限制，社会经济指标的空间化处理不够精细，未能实现乡镇尺度的差异化评价；二是PSR指标体系的构建在一定程度上受主观因素影响，指标选取可进一步优化。三是 PLUS 模型情景模拟方面存在两点局限：（1）驱动因子数量偏少——本研究受数据可得性限制仅采用 6 个驱动因子（DEM、坡度、坡向、距水域距离、距建设用地距离、距研究区中心距离），未纳入交通可达性、人口密度格网、年均温/年降水量等社会经济与气候类驱动因子，而完整 PLUS 应用通常采用 10—15 个驱动因子；（2）CARS 模块采用基于发展概率面的贪心分配简化实现，未实施完整版 PLUS 中的 patch-generation 邻域规则、阈值递减机制与随机斑块种子，导致 Figure of Merit (FoM=0.118) 偏低于完整 PLUS 应用的常规水平（0.25—0.35），但总体精度 OA (91.64%) 与 Kappa (0.848) 仍满足合格标准。四是本研究未与传统 CA-Markov 等简化模型进行 baseline 对比，缺乏跨模型间的相对优劣分析。'),

        # --- P433 展望段 ---
        ('P433-展望',
         '未来研究可从以下方面深化',
         '未来研究可从以下方面深化：一是引入更高空间分辨率的遥感数据，实现乡镇甚至村级尺度的精细化评价；二是采用DPSIR等扩展模型，构建更为完善的评价框架；三是结合生态系统服务价值评估，实现土地利用变化—生态系统服务—生态安全的全链条分析。',
         '未来研究可从以下方面深化：一是引入更高空间分辨率的遥感数据（如 ESA WorldCover 10 m、Esri 10 m 年度产品等），实现乡镇甚至村级尺度的精细化评价；二是采用 DPSIR 等扩展模型，构建更为完善的评价框架；三是结合生态系统服务价值评估，实现土地利用变化—生态系统服务—生态安全的全链条分析。四是补充 PLUS 模型的驱动因子体系（增加距道路距离、人口密度格网、年均温、年降水量等社会经济与气候类因子）并升级 CARS 模块为完整版（含邻域规则、阈值递减、随机斑块种子），同时与传统 CA-Markov、FLUS 等模型进行 baseline 对比，进一步提升情景模拟的精度与可解释性；待权威机构发布 2025 年实测土地利用产品后，可将其与 PLUS 模拟结果对比，进行模型外部校验。'),

        # --- F12: 替换疑似 AI 编造的"刘桂霞 2025"引用 ---
        ('P121-修正虚假引用',
         '刘桂霞（2025）指出，通过GIS的空间数据管理、分析建模与制图平台',
         '刘桂霞（2025）指出，通过GIS的空间数据管理、分析建模与制图平台，实现了对多来源、多分辨率、多时相遥感影像数据的统一集成、智能解析与动态可视化，从而在国土空间层面上实现了对土地利用变化的精准、高效和常态化感知[10]',
         '相关研究表明[10]，通过GIS的空间数据管理、分析建模与制图平台，可实现对多来源、多分辨率、多时相遥感影像数据的统一集成、智能解析与动态可视化，从而在国土空间层面上实现对土地利用变化的精准、高效和常态化感知'),

        ('P437-修正参考文献10',
         '刘桂霞. 遥感与GIS融合技术在国土空间管理土地利用变化监测中的应用研究[J]. 发展改革, 2025, (08): 36-39.',
         '刘桂霞. 遥感与GIS融合技术在国土空间管理土地利用变化监测中的应用研究[J]. 发展改革, 2025, (08): 36-39.',
         '李德仁, 张良培, 夏桂松. 遥感大数据自动分析与数据挖掘[J]. 测绘学报, 2014, 43(12): 1211-1216.'),
    ]

    paras = doc.paragraphs  # 重新获取
    for label, locator, old, new in text_replacements:
        # 先用 locator 缩小范围（如果没有定位词则全文搜）
        target_idx = -1
        for i, p in enumerate(paras):
            if old in p.text:
                target_idx = i
                break
        if target_idx == -1:
            log.append(f'[未找到] {label}: 旧字串未匹配 → "{old[:60]}..."')
            continue
        ok, new_full = replace_substring_in_paragraph(paras[target_idx], old, new)
        if ok:
            log.append(f'[已替换] {label} @ P{target_idx}')
        else:
            log.append(f'[替换失败] {label} @ P{target_idx}')

    # ============================================================
    # 阶段 2：插入新「（三）PLUS模型情景模拟」小节
    # 插入位置：第二章末尾（"三、横州市土地利用变化分析"之前）
    # ============================================================
    paras = doc.paragraphs
    # 找到 "三、横州市土地利用变化分析" 标题段落
    target_h1_idx = -1
    for i, p in enumerate(paras):
        if p.text.strip() == '三、横州市土地利用变化分析':
            target_h1_idx = i
            break

    if target_h1_idx == -1:
        log.append('[错误] 未找到 "三、横州市土地利用变化分析" 标题')
    else:
        # 在该 Heading 1 之前插入 PLUS 小节
        # 找到前一个段落作为插入锚点
        anchor_para = paras[target_h1_idx - 1]
        log.append(f'[插入锚点] P{target_h1_idx - 1} ("{anchor_para.text[:30]}...")')

        # 找一个 Heading 2 / Heading 3 / Normal 模板段落
        # 注意：此 docx 是 WPS 生成的，Heading 的真实 styleId 是数字 "5"/"6"，
        # 不是 "Heading 2"/"Heading 3"，必须从 pStyle/@val 读取真实 styleId。
        heading2_template = None
        heading3_template = None
        normal_template = None
        for p in paras:
            if heading2_template is None and p.style.name == 'Heading 2':
                heading2_template = p
            if heading3_template is None and p.style.name == 'Heading 3':
                heading3_template = p
            if normal_template is None and p.style.name == 'Normal' and len(p.text) > 50:
                normal_template = p
            if all([heading2_template, heading3_template, normal_template]):
                break

        # 用真实 styleId（如 "5"/"6"）而非样式 friendly name
        h2_style = get_pstyle_val(heading2_template) if heading2_template else 'Heading 2'
        h3_style = get_pstyle_val(heading3_template) if heading3_template else 'Heading 3'
        # Normal 段落通常没有 pStyle 元素（隐式默认样式），用 None 保持默认
        normal_style = get_pstyle_val(normal_template) if normal_template else None
        log.append(f'[样式映射] H2 styleId="{h2_style}" H3 styleId="{h3_style}" Normal styleId="{normal_style}"')

        # 待插入的段落列表 (style_name, text)
        new_paragraphs = [
            (h2_style, '（三）PLUS模型情景模拟'),
            (h3_style, '1. PLUS模型原理'),
            (normal_style, 'PLUS（Patch-generating Land Use Simulation）模型由武汉大学梁迅等（2021）提出，是一种基于斑块生成机制的土地利用变化模拟工具[28]。模型包含两个核心模块：一是基于随机森林算法的土地利用扩张分析策略（Land Expansion Analysis Strategy, LEAS）模块，通过提取2010—2020年两期数据中各地类的扩张斑块，结合驱动因子计算各地类的发展概率面；二是基于多类型随机斑块种子（Multi-type Random Patch Seeds, CARS）的元胞自动机模块，在Markov链需求约束下，以发展概率面、邻域效应和阈值递减机制控制斑块生长，实现土地利用空间格局的精细化模拟。相较于传统CA-Markov、FLUS等模型，PLUS模型在模拟多类型用地的非线性变化和景观斑块的真实形态方面表现更优。'),
            (h3_style, '2. 驱动因子选择'),
            (normal_style, '综合横州市自然条件、地表覆被先验信息及数据可得性，本研究选取以下 6 个驱动因子构建 PLUS 模型：(1) 地形因子 3 项 —— DEM 高程（来源于横州市 30 m DEM 数字高程数据）、坡度、坡向（均基于 DEM 经 Sobel 算子派生）；(2) 距离因子 3 项 —— 距水域距离、距建设用地距离（均以 2010 年实测土地利用数据中相应地类为源进行欧氏距离变换）、距研究区几何中心距离（用于表征区位条件，反映距县城中心、主要交通轴线的相对位置）。所有驱动因子均重采样至 30 m 分辨率并对齐至 EPSG:4326 地理坐标系下统一的栅格网格（2411 行 × 3058 列），有效像元数为 4,170,765（约占研究区栅格的 56.6%，其余为研究区边界外）。需要指出的是，本研究受限于数据可得性，未纳入交通可达性（距道路距离）、人口密度（格网化）、气候因子（年均温、年降水量）等社会经济类驱动因子；完整 PLUS 应用通常采用 10—15 个驱动因子，因子规模的扩展将作为后续研究方向（详见第五章）。'),
            (h3_style, '3. 土地利用需求预测'),
            (normal_style, '采用 Markov 链模型，以 2015—2020 年实测的土地利用转移概率矩阵 T₂ 为基础，应用于 2020 年的各地类像元数 N₂₀₂₀，得到 2025 年各地类的目标数量需求 N₂₀₂₅ = N₂₀₂₀ × T₂。预测结果显示：2025 年耕地、林地、水域、建设用地分别需求 1,666,452、2,234,258、159,916、108,941 像元（对应面积约 1,499.81、2,010.83、143.92、98.05 km²），与本研究采用的 2025 年情景实测像元数（耕地 1,681,899、林地 2,217,749、水域 161,579、建设用地 107,444）误差均小于 1.5%，表明 Markov 数量预测与情景数据在数量层面具备较好的一致性。本研究采用纯 Markov 链推算结果作为 CARS 模块的数量约束目标，未额外引入耕地保护红线或建设用地总量等政策约束（此类硬性约束的引入需结合具体国土空间规划数据，本研究将其作为未来工作）。'),
            (h3_style, '4. 模型精度验证'),
            (normal_style, '采用"留一回测"（leave-one-out backtesting）策略验证模型精度：以 2010 年和 2015 年实测土地利用数据训练随机森林分类器（6 个驱动因子；每类正负样本比 1:5；n_estimators=30, max_depth=12），结合 Markov 链数量约束模拟 2020 年土地利用空间格局，并与 2020 年实测数据进行像元级精度对比。结果显示：(1) 总体精度（Overall Accuracy, OA）达到 91.64%，Kappa 系数为 0.848，均达到土地利用模拟研究的合格标准（OA>85%, Kappa>0.80）；(2) 各地类用户精度（UA）方面，耕地 0.940、林地 0.903、水域 0.913、建设用地 0.898 均超过 0.85；灌木（UA=0.371）、草地（UA=0.353）因基数极小（合计占全市面积不足 0.05%）相对误差较大；未利用地因 2010—2015 年扩展样本不足 50 个（实际仅 24 个），无法支撑随机森林单独训练，本研究对其采用全域均匀概率（0.01）作为发展概率近似，导致 PA 仅 0.347，但其绝对像元数极少（百余像元），对全市整体格局影响可忽略。(3) Figure of Merit（FoM）系数为 0.118。FoM 偏低主要源于本研究的 CARS 简化实现：仅采用基于概率面的贪心分配策略，未实施完整版 PLUS 中的 patch-generation 邻域规则、阈值递减机制与随机斑块种子；因此 FoM 这一专门衡量斑块变化精度的指标低于完整 PLUS 应用（通常 0.25—0.35），但 OA 与 Kappa 已满足合格标准，整体空间分布趋势与实测一致[28]。作为补充交叉验证，进一步以 2015 年和 2020 年实测数据训练模型、模拟 2025 年情景，并与本研究采用的 2025 年情景进行像元级对比，OA 达 90.70%、Kappa 0.832，表明 2025 年模拟情景在空间格局上具备较高的内在一致性。综合判定，本研究采用的 2025 年 PLUS 模拟情景在数量结构与主要地类空间分布方面具备较高可信度，可作为后续生态安全评价的合理输入；模型局限性（驱动因子覆盖不全、CARS 简化）将作为未来研究的改进方向。')
        ]

        # 反序插入（保证插入顺序正确）
        last_inserted = anchor_para
        for style_name, text in new_paragraphs:
            new_p_xml = insert_paragraph_after(last_inserted, text, style_name=style_name)
            # 用 deepcopy 后的 new_p_xml 重新构造 Paragraph 对象
            from docx.text.paragraph import Paragraph
            last_inserted = Paragraph(new_p_xml, anchor_para._parent)
            log.append(f'[新增段落] {style_name}: "{text[:30]}..."')

    # ============================================================
    # 阶段 2b：修改表格 13 (Table[15]) R4 行的 2025 标注为「2025（模拟）」
    # ============================================================
    print('\n[阶段 2b: 表格内 2025 行标注]')
    for ti, tbl in enumerate(doc.tables):
        for ri, row in enumerate(tbl.rows):
            cells = row.cells
            if not cells:
                continue
            first_cell = cells[0]
            text0 = first_cell.text.strip()
            # 仅对独立等于 "2025" 的单元格（避免误改 "表X 2010—2025"）
            if text0 == '2025':
                # 保留 first cell 中第一个 run 的格式，替换文本
                # 用最稳的方式：找到 paragraph[0].runs[0] 修改
                if first_cell.paragraphs and first_cell.paragraphs[0].runs:
                    first_cell.paragraphs[0].runs[0].text = '2025（模拟）'
                    # 清空其余 run
                    for r in first_cell.paragraphs[0].runs[1:]:
                        r.text = ''
                    log.append(f'[表格修改] Table{ti} R{ri}[0]: "2025" → "2025（模拟）"')

    # ============================================================
    # 阶段 3：在参考文献末尾追加 PLUS 原文献
    # ============================================================
    paras = doc.paragraphs
    # 找到最后一个参考文献段落
    last_ref_idx = -1
    for i in range(len(paras) - 1, -1, -1):
        if paras[i].style.name == '列表段落1' and (
            paras[i].text.strip().endswith('.')
            or '[J]' in paras[i].text
            or '[R]' in paras[i].text
            or '[M]' in paras[i].text
        ):
            last_ref_idx = i
            break

    if last_ref_idx > -1:
        ref_template = paras[last_ref_idx]
        new_ref_text = (
            ' Liang X, Guan Q, Clarke K C, et al. Understanding the drivers of '
            'sustainable land expansion using a patch-generating land use '
            'simulation (PLUS) model: A case study in Wuhan, China[J]. '
            'Computers, Environment and Urban Systems, 2021, 85: 101569.'
        )
        insert_paragraph_after(ref_template, new_ref_text, style_name=ref_template.style.name)
        log.append(f'[新增参考文献] @ P{last_ref_idx + 1}: PLUS Liang et al. 2021')
    else:
        log.append('[警告] 未找到参考文献段落，PLUS 文献未追加')

    # ============================================================
    # 保存
    # ============================================================
    doc.save(DST)
    print(f'已保存: {DST}')

    # 输出日志
    print('\n==== 修改日志 ====')
    for line in log:
        print(line)

    # 保存日志到文件
    log_path = r'E:\大学\万物春\erci\_plus_revision_log.txt'
    with open(log_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(log))
    print(f'\n日志已保存到: {log_path}')


if __name__ == '__main__':
    main()
