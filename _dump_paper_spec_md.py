# -*- coding: utf-8 -*-
"""把附件三论文撰写规范 docx 全文导出为 markdown 文本（UTF-8）"""
from docx import Document

SPEC = r'E:\大学\万物春\erci\附件三：论文撰写规范 (1).docx'
OUT = r'f:\Gorsachius magnificus\_paper_spec_raw.md'

doc = Document(SPEC)
with open(OUT, 'w', encoding='utf-8') as f:
    f.write('# 附件三：论文撰写规范 - 全文导出\n\n')
    f.write(f'段落数: {len(doc.paragraphs)} | 表格数: {len(doc.tables)}\n\n')
    f.write('---\n\n')
    for i, p in enumerate(doc.paragraphs):
        style = p.style.name if p.style else 'None'
        text = p.text
        if not text.strip():
            continue
        f.write(f'[P{i:3d}|{style}] {text}\n\n')
    if doc.tables:
        f.write('\n## 表格\n\n')
        for ti, tbl in enumerate(doc.tables):
            f.write(f'\n### Table {ti} ({len(tbl.rows)} rows × {len(tbl.columns)} cols)\n\n')
            for ri, row in enumerate(tbl.rows):
                for ci, cell in enumerate(row.cells):
                    txt = cell.text.replace('\n', ' | ')
                    f.write(f'  R{ri}C{ci}: {txt}\n')

print(f'已导出至: {OUT}')
