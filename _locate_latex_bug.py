# -*- coding: utf-8 -*-
"""定位 docx 中未渲染的 LaTeX 源码"""
import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document

V3 = r'E:\大学\万物春\erci\郑春铃+横州市土地利用变化与生态安全评价(1)_PLUS修订版_v3.docx'
ORIG = r'E:\大学\万物春\erci\郑春铃+横州市土地利用变化与生态安全评价(1).docx.bak_before_plus_revision'

print('=== v3 中含 LaTeX 源码的段落 ===')
doc = Document(V3)
hits = []
for i, p in enumerate(doc.paragraphs):
    if 'x_{ij}' in p.text or 'x_{i' in p.text or 'min(x_j)' in p.text or 'max(x_j)' in p.text or '\\prime' in p.text or '\\frac' in p.text:
        hits.append((i, p.style.name, p.text))
        print(f'\nP{i} | style={p.style.name}')
        print(f'  text: {p.text[:500]}')

print(f'\n共 {len(hits)} 处命中')

# 同步检查原文档（无修订）中同样位置
print('\n\n=== 原文档（bak）中含 LaTeX 源码的段落 ===')
doc2 = Document(ORIG)
for i, p in enumerate(doc2.paragraphs):
    if 'x_{ij}' in p.text or 'x_{i' in p.text or 'min(x_j)' in p.text or 'max(x_j)' in p.text or '\\prime' in p.text or '\\frac' in p.text:
        print(f'\nP{i} | style={p.style.name}')
        print(f'  text: {p.text[:500]}')

# 检查公式相关（OMML 公式）
print('\n\n=== v3 OMML 公式检查 ===')
from docx.oxml.ns import qn
for i, p in enumerate(doc.paragraphs):
    omml_count = len(p._p.findall('.//' + qn('m:oMath'))) + len(p._p.findall('.//' + qn('m:oMathPara')))
    if omml_count > 0:
        print(f'P{i}: {omml_count} OMML 公式 | text: {p.text[:100]}')
